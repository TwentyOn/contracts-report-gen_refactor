import io
import os
import psycopg2
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from dotenv import load_dotenv

from utils.postprocessing_report_file import upload_to_s3, write_s3path_to_bd

# Загружаем переменные из .env файла
load_dotenv('.env')


def get_month_name(month_num):
    """Преобразование номера месяца в название"""
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    return months.get(month_num, "месяца")


def format_number(number):
    """Форматирование числа с разделителями разрядов и десятичной частью"""
    if number is None:
        return "Не указано"

    try:
        # Преобразуем в float для обработки
        num = float(number)

        # Форматируем число с 2 знаками после запятой
        formatted = f"{num:,.2f}"

        # Заменяем запятую на пробел для разделения разрядов и точку на запятую для десятичной части
        formatted = formatted.replace(",", " ").replace(".", ",")

        return formatted

    except (ValueError, TypeError):
        return str(number) if number else "Не указано"


def format_amount_rubles_kopecks(amount):
    """Форматирование суммы в формате 'X рублей YY копеек'"""
    if amount is None:
        return "0 рублей 00 копеек"

    try:
        # Преобразуем в float
        num = float(amount)

        # Разделяем на рубли и копейки, используя округление для избежания ошибок с плавающей точкой
        rubles = int(num)
        kopecks = round((num - rubles) * 100)

        # Если копейки получились 100 из-за округления, корректируем
        if kopecks >= 100:
            rubles += 1
            kopecks = 0

        return f"{rubles} рублей {kopecks:02d} копеек"

    except (ValueError, TypeError):
        return "0 рублей 00 копеек"


# ID отчета для обработки (задавайте вручную)
# REPORT_ID = 15

# Параметры подключения к БД
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'database': os.getenv('DB_NAME', 'your_database'),
    'user': os.getenv('DB_USER', 'your_user'),
    'password': os.getenv('DB_PASSWORD', 'your_password'),
    'port': os.getenv('DB_PORT', '5432')
}


def connect_to_db():
    """Подключение к базе данных с обработкой ошибок"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"Ошибка подключения к БД: {e}")
        return None


def get_act_data(conn, report_id):
    """Получение данных для акта"""
    try:
        cursor = conn.cursor()

        # Сначала проверим, есть ли отчет вообще
        check_query = "SELECT id FROM gen_report_context_contracts.reports WHERE id = %s"
        cursor.execute(check_query, (report_id,))
        report_exists = cursor.fetchone()

        if not report_exists:
            print(f"Отчет с ID {report_id} не найден в таблице reports")
            return None

        print(f"Отчет {report_id} найден, проверяем связи...")

        # Получим данные отчета
        simple_query = "SELECT id, id_contracts, id_requests FROM gen_report_context_contracts.reports WHERE id = %s"
        cursor.execute(simple_query, (report_id,))
        simple_result = cursor.fetchone()

        if not simple_result:
            print("Не удалось получить данные отчета")
            return None

        print(f"Данные отчета: ID={simple_result[0]}, ID контракта={simple_result[1]}, ID заявки={simple_result[2]}")

        # Проверим контракт
        contract_query = "SELECT id, id_customer, id_contractor, date_contract, number_contract, theme_contract, service_name FROM gen_report_context_contracts.contracts WHERE id = %s"
        cursor.execute(contract_query, (simple_result[1],))
        contract_result = cursor.fetchone()

        if contract_result:
            print(
                f"Контракт найден: ID={contract_result[0]}, ID заказчика={contract_result[1]}, ID исполнителя={contract_result[2]}")
            print(f"Дата контракта: {contract_result[3]}, Номер: {contract_result[4]}, Тема: {contract_result[5]}")

            # Получаем тексты из textforformdocument по ключам act_1-act_27
            text_queries = [
                ("act_1", "act_1_text"),
                ("act_2", "act_2_text"),
                ("act_3", "act_3_text"),
                ("act_4", "act_4_text"),
                ("act_5", "act_5_text"),
                ("act_6", "act_6_text"),
                ("act_7", "act_7_text"),
                ("act_8", "act_8_text"),
                ("act_9", "act_9_text"),
                ("act_10", "act_10_text"),
                ("act_11", "act_11_text"),
                ("act_12", "act_12_text"),
                ("act_13", "act_13_text"),
                ("act_14", "act_14_text"),
                ("act_15", "act_15_text"),
                ("act_16", "act_16_text"),
                ("act_17", "act_17_text"),
                ("act_18", "act_18_text"),
                ("act_19", "act_19_text"),
                ("act_20", "act_20_text"),
                ("act_21", "act_21_text"),
                ("act_22", "act_22_text"),
                ("act_23", "act_23_text"),
                ("act_24", "act_24_text"),
                ("act_25", "act_25_text"),
                ("act_26", "act_26_text"),
                ("act_27", "act_27_text"),
                ("act_28", "act_28_text"),
                ("act_29", "act_29_text")
            ]

            texts = {}
            for key, var_name in text_queries:
                text_query = f"SELECT text_data FROM gen_report_context_contracts.textforformdocument WHERE key = '{key}'"
                cursor.execute(text_query)
                text_result = cursor.fetchone()
                texts[var_name] = text_result[0] if text_result else f"Текст {key} не найден"
                print(f"Текст {key}: {texts[var_name]}")

            # Получаем данные о заказчике (id_customer)
            customer_query = "SELECT full_name_nominative, position_genitive, representative_name_genitive, representative_basis, long_name_organisation, position_nominative, representative_signature FROM gen_report_context_contracts.organizations WHERE id = %s"
            cursor.execute(customer_query, (contract_result[1],))
            customer_result = cursor.fetchone()

            if customer_result:
                print(f"Заказчик найден: {customer_result[0]}")
            else:
                print(f"Заказчик с ID {contract_result[1]} не найден")
                customer_result = (None, None, None, None, None, None, None)

            # Получаем данные об исполнителе (id_contractor)
            contractor_query = "SELECT full_name_nominative, short_name_organisation, position_genitive, representative_name_genitive, representative_basis, long_name_organisation, position_nominative, representative_signature FROM gen_report_context_contracts.organizations WHERE id = %s"
            cursor.execute(contractor_query, (contract_result[2],))
            contractor_result = cursor.fetchone()

            if contractor_result:
                print(f"Исполнитель найден: {contractor_result[0]}")
            else:
                print(f"Исполнитель с ID {contract_result[2]} не найден")
                contractor_result = (None, None, None, None, None, None, None, None)

            # Получаем данные из таблицы requests по id_requests из отчета
            requests_query = "SELECT start_date, end_date, application_number, date_request, financial_unit, financial_quantity, financial_price_per_unit, financial_total_amount, financial_quality, financial_vat_amount, financial_vat_amount_words, advance_payment_transferred, advance_payment_credited, amount_due, amount_due_words FROM gen_report_context_contracts.requests WHERE id = %s"
            cursor.execute(requests_query, (simple_result[2],))
            requests_result = cursor.fetchone()

            if requests_result:
                print(
                    f"Заявка найдена: start_date={requests_result[0]}, end_date={requests_result[1]}, application_number={requests_result[2]}, date_request={requests_result[3]}")
                print(
                    f"Финансовые данные: unit={requests_result[4]}, quantity={requests_result[5]}, price={requests_result[6]}, total={requests_result[7]}, quality={requests_result[8]}")
                print(
                    f"Суммы прописью: total_words={requests_result[9]}, vat={requests_result[10]}, vat_words={requests_result[11]}")
            else:
                print(f"Заявка для контракта {contract_result[0]} не найдена")
                requests_result = (None, None, None, None, None, None, None, None, None, None, None, None)

            # Номер договора уже получен из таблицы contracts
            print(f"Организация найдена, номер договора: {contract_result[4]}")

            return {
                'report_id': simple_result[0],
                'contract_id': contract_result[0],
                'customer_id': contract_result[1],
                'contractor_id': contract_result[2],
                'date_contract': contract_result[3],
                'number_contract': contract_result[4],
                'theme_contract': contract_result[5],
                'service_name': contract_result[6],
                'customer_full_name_nominative': customer_result[0],
                'customer_position_genitive': customer_result[1],
                'customer_representative_name_genitive': customer_result[2],
                'customer_representative_basis': customer_result[3],
                'customer_long_name_organisation': customer_result[4],
                'customer_position_nominative': customer_result[5],
                'customer_representative_signature': customer_result[6],
                'contractor_full_name_nominative': contractor_result[0],
                'contractor_short_name_organisation': contractor_result[1],
                'contractor_position_genitive': contractor_result[2],
                'contractor_representative_name_genitive': contractor_result[3],
                'contractor_representative_basis': contractor_result[4],
                'contractor_long_name_organisation': contractor_result[5],
                'contractor_position_nominative': contractor_result[6],
                'contractor_representative_signature': contractor_result[7],
                'start_date': requests_result[0],
                'end_date': requests_result[1],
                'application_number': requests_result[2],
                'date_request': requests_result[3],
                'financial_unit': requests_result[4],
                'financial_quantity': requests_result[5],
                'financial_price_per_unit': requests_result[6],
                'financial_total_amount': requests_result[7],
                'financial_quality': requests_result[8],
                'financial_vat_amount': requests_result[9],
                'financial_vat_amount_words': requests_result[10],
                'advance_payment_transferred': requests_result[11],
                'advance_payment_credited': requests_result[12],
                'amount_due': requests_result[13],
                'amount_due_words': requests_result[14],
                **texts
            }
        else:
            print(f"Контракт с ID {simple_result[1]} не найден")

        return None

    except Exception as e:
        print(f"Ошибка при получении данных: {e}")
        return None
    finally:
        cursor.close()


def create_act_document(data, report_id):
    """Создание акта в формате Word"""
    try:
        print(f"Создаем акт с данными: {data}")

        doc = Document()

        # 1. Добавляем заголовок с центрированием и жирным шрифтом
        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Формируем дату в нужном формате
        if data['date_contract']:
            contract_date = data['date_contract']
            if isinstance(contract_date, str):
                contract_date = datetime.strptime(contract_date, '%Y-%m-%d').date()
            day = contract_date.day
            month_name = get_month_name(contract_date.month)
            year = contract_date.year
            date_str = f"от {day} «{month_name}» {year} г."
        else:
            date_str = "дата не указана"

        # Формируем заголовок
        header_text = f"{data['act_1_text']} {date_str} № {data['number_contract']} {data['theme_contract']}"

        header_run = header_paragraph.add_run(header_text)
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(11)
        header_run.bold = True
        print(f"Добавлен заголовок: '{header_text}'")

        # 2. Добавляем пустую строку
        doc.add_paragraph()

        # 3. Создаем таблицу 2x2
        table = doc.add_table(rows=2, cols=2)

        # Убираем видимые границы таблицы
        for row in table.rows:
            for cell in row.cells:
                # Убираем все границы
                cell._tc.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._tc.get_or_add_tcPr().find(qn('w:tcBorders'))
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    borders.append(border)

                # Настраиваем выравнивание по центру и середине
                cell.vertical_alignment = 1  # По середине
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт для текста в ячейках
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

        # Устанавливаем ширину таблицы на всю страницу
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3.25)  # Половина ширины страницы A4

        # Заполняем ячейки с правильным форматированием
        # Левая верхняя ячейка = act_2
        cell_00 = table.cell(0, 0)
        cell_00.text = data['act_2_text']
        for paragraph in cell_00.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # Правая верхняя ячейка = act_3
        cell_01 = table.cell(0, 1)
        cell_01.text = data['act_3_text']
        for paragraph in cell_01.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # Левая нижняя ячейка = act_4
        cell_10 = table.cell(1, 0)
        cell_10.text = data['act_4_text']
        for paragraph in cell_10.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # Правая нижняя ячейка = текущая дата в формате "«04» июля 2025 г., Москва"
        current_date = datetime.now()
        day = current_date.day
        month_name = get_month_name(current_date.month)
        year = current_date.year
        date_text = f"«{day:02d}» {month_name} {year} г., Москва"
        cell_11 = table.cell(1, 1)
        cell_11.text = date_text
        for paragraph in cell_11.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        print(f"Добавлена таблица с данными:")
        print(f"  act_2: {data['act_2_text']}")
        print(f"  act_3: {data['act_3_text']}")
        print(f"  act_4: {data['act_4_text']}")
        print(f"  Дата: {date_text}")

        # 4. Добавляем пустую строку после таблицы
        doc.add_paragraph()

        # 5. Добавляем большое предложение с абзацным отступом
        big_paragraph = doc.add_paragraph()
        big_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        big_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        big_paragraph.paragraph_format.line_spacing = 1.5

        # Формируем большое предложение с правильным форматированием
        # Часть о заказчике
        if data['customer_full_name_nominative']:
            run = big_paragraph.add_run(data['customer_full_name_nominative'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            big_paragraph.add_run(" ")  # Всегда добавляем пробел после данных из БД

        run = big_paragraph.add_run(data['act_5_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        run = big_paragraph.add_run(data['act_6_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        # Пробелы в textforformdocument регулируются в БД

        # Добавляем act_11 после act_6
        run = big_paragraph.add_run(data['act_11_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        if data['customer_position_genitive']:
            run = big_paragraph.add_run(data['customer_position_genitive'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            big_paragraph.add_run(" ")  # Всегда добавляем пробел после данных из БД

        if data['customer_representative_name_genitive']:
            run = big_paragraph.add_run(data['customer_representative_name_genitive'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Не добавляем пробел после representative_name_genitive

        run = big_paragraph.add_run(data['act_7_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        if data['customer_representative_basis']:
            run = big_paragraph.add_run(data['customer_representative_basis'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Не добавляем пробел после representative_basis

        run = big_paragraph.add_run(data['act_8_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        # Часть об исполнителе
        if data['contractor_full_name_nominative']:
            # Делаем первую букву маленькой для середины предложения
            contractor_name = data['contractor_full_name_nominative']
            if contractor_name:
                contractor_name = contractor_name[0].lower() + contractor_name[1:]
            run = big_paragraph.add_run(contractor_name)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            big_paragraph.add_run(" ")  # Всегда добавляем пробел после данных из БД

        if data['contractor_short_name_organisation']:
            run = big_paragraph.add_run("(")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run = big_paragraph.add_run(data['contractor_short_name_organisation'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            run = big_paragraph.add_run(")")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Не добавляем пробел после скобок с short_name_organisation

        run = big_paragraph.add_run(data['act_9_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        if data['contractor_position_genitive']:
            run = big_paragraph.add_run(data['contractor_position_genitive'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            big_paragraph.add_run(" ")  # Всегда добавляем пробел после данных из БД

        if data['contractor_representative_name_genitive']:
            run = big_paragraph.add_run(data['contractor_representative_name_genitive'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Не добавляем пробел после representative_name_genitive

        run = big_paragraph.add_run(data['act_7_text'])  # Снова act_7
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        if data['contractor_representative_basis']:
            run = big_paragraph.add_run(data['contractor_representative_basis'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Не добавляем пробел после representative_basis

        run = big_paragraph.add_run(data['act_10_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        print(f"Добавлено большое предложение с данными заказчика и исполнителя")

        # 6. Добавляем второй абзац с тем же отступом
        second_paragraph = doc.add_paragraph()
        second_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        second_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        second_paragraph.paragraph_format.line_spacing = 1.5

        # 1. act_12
        run = second_paragraph.add_run(data['act_12_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД
        second_paragraph.add_run(" ")  # Пробел после act_12

        # 2. Дата в формате "«28» декабря 2024 г. №500-976"
        if data['date_contract']:
            contract_date = data['date_contract']
            if isinstance(contract_date, str):
                contract_date = datetime.strptime(contract_date, '%Y-%m-%d').date()
            day = contract_date.day
            month_name = get_month_name(contract_date.month)
            year = contract_date.year
            date_text = f"«{day}» {month_name} {year} г. №{data['number_contract']}"
        else:
            date_text = "дата не указана"

        run = second_paragraph.add_run(date_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        second_paragraph.add_run(" ")  # Пробел после даты

        # 3. theme_contract
        run = second_paragraph.add_run(data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        second_paragraph.add_run(" ")  # Пробел после темы

        # 4. Статический текст
        static_text = "(далее – Договор) Исполнителем оказаны услуги, а Заказчиком приняты услуги в рамках Договора."
        run = second_paragraph.add_run(static_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        print(f"Добавлен второй абзац с act_12, датой, темой и статическим текстом")

        # 7. Добавляем третий абзац с тем же отступом
        third_paragraph = doc.add_paragraph()
        third_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        third_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        third_paragraph.paragraph_format.line_spacing = 1.5

        # 1. act_13
        run = third_paragraph.add_run(data['act_13_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД
        third_paragraph.add_run(" ")  # Пробел после act_13

        # 2. Период с start_date по end_date
        if data['start_date'] and data['end_date']:
            start_date = data['start_date']
            end_date = data['end_date']

            if isinstance(start_date, str):
                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
            if isinstance(end_date, str):
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

            start_day = start_date.day
            start_month_name = get_month_name(start_date.month)
            start_year = start_date.year

            end_day = end_date.day
            end_month_name = get_month_name(end_date.month)
            end_year = end_date.year

            period_text = f"с «{start_day}» {start_month_name} {start_year} г. по «{end_day}» {end_month_name} {end_year} г."
        else:
            period_text = "период не указан"

        run = third_paragraph.add_run(period_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        third_paragraph.add_run(" ")  # Пробел после периода

        # 3. act_14 + номер заявки
        run = third_paragraph.add_run(data['act_14_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        if data['application_number']:
            run = third_paragraph.add_run(data['application_number'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            third_paragraph.add_run(" ")  # Пробел после номера заявки

        # 4. Дата заявки в формате "от «24» марта 2025 года:"
        if data['date_request']:
            request_date = data['date_request']
            if isinstance(request_date, str):
                request_date = datetime.strptime(request_date, '%Y-%m-%d').date()

            req_day = request_date.day
            req_month_name = get_month_name(request_date.month)
            req_year = request_date.year

            request_date_text = f"от «{req_day}» {req_month_name} {req_year} года:"
        else:
            request_date_text = "дата заявки не указана:"

        run = third_paragraph.add_run(request_date_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        print(f"Добавлен третий абзац с act_13, периодом, act_14+номер заявки и датой заявки")

        # 8. Добавляем таблицу 6x2
        data_table = doc.add_table(rows=2, cols=6)
        data_table.style = 'Table Grid'  # Видимые границы

        # Настраиваем выравнивание ячеек по центру и середине
        for row in data_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1  # По середине
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт для текста в ячейках
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

        # Устанавливаем ширину таблицы на всю страницу
        data_table.autofit = False
        data_table.allow_autofit = False
        for row in data_table.rows:
            for cell in row.cells:
                cell.width = Inches(1.1)  # Равномерное распределение по 6 колонкам

        # Заполняем первую строку (заголовки)
        headers = [
            data['act_15_text'],
            data['act_16_text'],
            data['act_17_text'],
            data['act_18_text'],
            data['act_19_text'],
            data['act_20_text']
        ]

        for i, header in enumerate(headers):
            cell = data_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

        # Заполняем вторую строку (данные)
        data_values = [
            data['service_name'] or "Не указано",
            data['financial_unit'] or "Не указано",
            str(data['financial_quantity']) if data['financial_quantity'] else "Не указано",
            format_number(data['financial_price_per_unit']),
            format_number(data['financial_total_amount']),
            data['financial_quality'] or "Не указано"
        ]

        for i, value in enumerate(data_values):
            cell = data_table.cell(1, i)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

        print(f"Добавлена таблица 6x2 с данными:")
        print(f"  Заголовки: {headers}")
        print(f"  Данные: {data_values}")

        # 9. Добавляем пустую строку
        doc.add_paragraph()

        # 10. Добавляем абзац с act_28 + сумма аванса переведенного
        advance_transferred_paragraph = doc.add_paragraph()
        advance_transferred_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        advance_transferred_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        advance_transferred_paragraph.paragraph_format.line_spacing = 1.5

        # act_28
        run = advance_transferred_paragraph.add_run(data['act_28_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        advance_transferred_paragraph.add_run(" ")

        # Сумма аванса переведенного
        advance_amount = format_amount_rubles_kopecks(data['advance_payment_transferred'])
        run = advance_transferred_paragraph.add_run(advance_amount)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        advance_transferred_paragraph.add_run(".")

        print(f"Добавлен абзац с act_28 и суммой аванса: {advance_amount}")

        # 11. Добавляем четвертый абзац (act_21)
        fourth_paragraph = doc.add_paragraph()
        fourth_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        fourth_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        fourth_paragraph.paragraph_format.line_spacing = 1.5

        # act_21
        run = fourth_paragraph.add_run(data['act_21_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        print(f"Добавлен четвертый абзац с act_21")

        # 12. Добавляем абзац с act_29 + сумма аванса зачтенного
        advance_credited_paragraph = doc.add_paragraph()
        advance_credited_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        advance_credited_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        advance_credited_paragraph.paragraph_format.line_spacing = 1.5

        # act_29
        run = advance_credited_paragraph.add_run(data['act_29_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        advance_credited_paragraph.add_run(" ")

        # Сумма аванса зачтенного
        credited_amount = format_amount_rubles_kopecks(data['advance_payment_credited'])
        run = advance_credited_paragraph.add_run(credited_amount)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        advance_credited_paragraph.add_run(".")

        print(f"Добавлен абзац с act_29 и суммой аванса зачтенного: {credited_amount}")

        # 13. Добавляем пятый абзац (act_22 + суммы)
        fifth_paragraph = doc.add_paragraph()
        fifth_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        fifth_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        fifth_paragraph.paragraph_format.line_spacing = 1.5

        # act_22
        run = fifth_paragraph.add_run(data['act_22_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД
        fifth_paragraph.add_run(" ")  # Пробел после act_22

        # amount_due (форматированное число)
        run = fifth_paragraph.add_run(format_number(data['amount_due']))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        # amount_due_words в скобках
        if data['amount_due_words']:
            run = fifth_paragraph.add_run(" (")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run = fifth_paragraph.add_run(data['amount_due_words'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run = fifth_paragraph.add_run(")")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

        # act_23
        run = fifth_paragraph.add_run(data['act_23_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД
        fifth_paragraph.add_run(" ")  # Пробел после act_23

        # financial_vat_amount (форматированное число)
        run = fifth_paragraph.add_run(format_number(data['financial_vat_amount']))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        # financial_vat_amount_words в скобках
        if data['financial_vat_amount_words']:
            run = fifth_paragraph.add_run(" (")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run = fifth_paragraph.add_run(data['financial_vat_amount_words'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run = fifth_paragraph.add_run(")")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

        # Точка
        run = fifth_paragraph.add_run(".")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        print(f"Добавлен пятый абзац с act_22, суммами и act_23")

        # 14. Добавляем шестой абзац (act_24)
        sixth_paragraph = doc.add_paragraph()
        sixth_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        sixth_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Устанавливаем интервал между строками 1.5
        sixth_paragraph.paragraph_format.line_spacing = 1.5

        # act_24
        run = sixth_paragraph.add_run(data['act_24_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # Пробелы в textforformdocument регулируются в БД

        print(f"Добавлен шестой абзац с act_24")

        # 15. Добавляем пустую строку
        doc.add_paragraph()

        # 16. Добавляем таблицу 2x2 для подписей
        signature_table = doc.add_table(rows=2, cols=2)

        # Убираем видимые границы таблицы
        for row in signature_table.rows:
            for cell in row.cells:
                # Убираем все границы
                cell._tc.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._tc.get_or_add_tcPr().find(qn('w:tcBorders'))
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    borders.append(border)

                # Настраиваем выравнивание по центру и середине
                cell.vertical_alignment = 1  # По середине
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Устанавливаем шрифт для текста в ячейках
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

        # Устанавливаем ширину таблицы на всю страницу
        signature_table.autofit = False
        signature_table.allow_autofit = False
        for row in signature_table.rows:
            for cell in row.cells:
                cell.width = Inches(3.25)  # Половина ширины страницы A4

        # Заполняем левую верхнюю ячейку (заказчик)
        cell_00 = signature_table.cell(0, 0)
        cell_00_paragraph = cell_00.paragraphs[0]
        cell_00_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по левому краю

        # act_25 (жирным)
        run = cell_00_paragraph.add_run(data['act_25_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        cell_00_paragraph.add_run("\n")

        # long_name_organisation (жирным)
        if data['customer_long_name_organisation']:
            run = cell_00_paragraph.add_run(data['customer_long_name_organisation'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            cell_00_paragraph.add_run("\n")

        # position_nominative (жирным)
        if data['customer_position_nominative']:
            run = cell_00_paragraph.add_run(data['customer_position_nominative'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True

        # Заполняем правую верхнюю ячейку (исполнитель)
        cell_01 = signature_table.cell(0, 1)
        cell_01_paragraph = cell_01.paragraphs[0]
        cell_01_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по левому краю

        # act_26 (жирным)
        run = cell_01_paragraph.add_run(data['act_26_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        cell_01_paragraph.add_run("\n")

        # long_name_organisation (жирным)
        if data['contractor_long_name_organisation']:
            run = cell_01_paragraph.add_run(data['contractor_long_name_organisation'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            cell_01_paragraph.add_run("\n")

        # position_nominative (жирным)
        if data['contractor_position_nominative']:
            run = cell_01_paragraph.add_run(data['contractor_position_nominative'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True

        # Заполняем левую нижнюю ячейку (заказчик)
        cell_10 = signature_table.cell(1, 0)
        cell_10_paragraph = cell_10.paragraphs[0]
        cell_10_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по левому краю

        # Подпись заказчика
        if data['customer_representative_signature']:
            # Добавляем подчеркивания
            run = cell_10_paragraph.add_run("___________________/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Добавляем подпись жирным
            run = cell_10_paragraph.add_run(data['customer_representative_signature'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            run = cell_10_paragraph.add_run("/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        else:
            run = cell_10_paragraph.add_run("___________________/________________/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        cell_10_paragraph.add_run("\n")

        # act_27
        run = cell_10_paragraph.add_run(data['act_27_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        # Заполняем правую нижнюю ячейку (исполнитель)
        cell_11 = signature_table.cell(1, 1)
        cell_11_paragraph = cell_11.paragraphs[0]
        cell_11_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по левому краю

        # Подпись исполнителя
        if data['contractor_representative_signature']:
            # Добавляем подчеркивания
            run = cell_11_paragraph.add_run("___________________/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            # Добавляем подпись жирным
            run = cell_11_paragraph.add_run(data['contractor_representative_signature'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            run = cell_11_paragraph.add_run("/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        else:
            run = cell_11_paragraph.add_run("___________________/________________/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        cell_11_paragraph.add_run("\n")

        # act_27
        run = cell_11_paragraph.add_run(data['act_27_text'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        print(f"Добавлена таблица подписей 2x2 с данными заказчика и исполнителя")

        # Сохраняем документ с уникальным именем
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # output_path = f"Результат/акт_отчет_{REPORT_ID}_{timestamp}.docx"
        output_path = f"{report_id}/акт_{timestamp}.docx"

        file = io.BytesIO()
        doc.save(file)
        file.seek(0)
        return file, output_path
        print(f"Акт создан.")

    except Exception as e:
        print(f"Ошибка при создании акта: {e}")


def generate_act(report_id):
    """Основная функция"""
    print(f"Обработка отчета ID: {report_id}")

    # Подключение к БД
    conn = connect_to_db()
    if not conn:
        return

    try:
        # Получение данных
        data = get_act_data(conn, report_id)
        if not data:
            return

        print(f"Найден отчет: {data['report_id']}")
        print(f"ID контракта: {data['contract_id']}")
        print(f"ID заказчика: {data['customer_id']}")

        # Создание акта
        file, file_name = create_act_document(data, report_id)

        # отправка в S3
        # s3_filepath = upload_to_s3(minio_client, file, file_name)

        # запись пути к файлу в БД
        # write_s3path_to_bd(report_id, 'file_act', s3_filepath)


    except Exception as e:
        print(f"Общая ошибка: {e}")
        raise e
    finally:
        conn.close()
    return file, file_name


if __name__ == "__main__":
    generate_act(19)
