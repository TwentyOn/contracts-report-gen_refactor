import io
import os
import psycopg2
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from dotenv import load_dotenv
from minio import Minio

from utils.postprocessing_report_file import upload_to_s3
from utils.postprocessing_report_file import write_s3path_to_bd

# Загружаем переменные из .env файла
load_dotenv()


def get_month_name(month_num):
    """Преобразование номера месяца в название"""
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    return months.get(month_num, "месяца")


# ID отчета для обработки (задавайте вручную)
# REPORT_ID = 15

# Параметры подключения к БД
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'database': os.getenv('DB_NAME', 'your_database'),
    'user': os.getenv('DB_USER', 'your_user'),
    'password': os.getenv('DB_PASSWORD', 'your_password'),
    'port': os.getenv('DB_PORT', '5432')
}


def connect_to_db():
    """Подключение к базе данных с обработкой ошибок"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"Ошибка подключения к БД: {e}")
        return None


def get_report_data(conn, report_id):
    """Получение данных отчета, контракта и организации"""
    try:
        cursor = conn.cursor()

        # Сначала проверим, есть ли отчет вообще
        check_query = "SELECT id FROM gen_report_context_contracts.reports WHERE id = %s"
        cursor.execute(check_query, (report_id,))
        report_exists = cursor.fetchone()

        if not report_exists:
            print(f"Отчет с ID {report_id} не найден в таблице reports")
            return None

        print(f"Отчет {report_id} найден, проверяем связи...")

        # Получим данные отчета
        simple_query = "SELECT id, id_contracts, id_requests FROM gen_report_context_contracts.reports WHERE id = %s"
        cursor.execute(simple_query, (report_id,))
        simple_result = cursor.fetchone()

        if not simple_result:
            print("Не удалось получить данные отчета")
            return None

        print(f"Данные отчета: ID={simple_result[0]}, ID контракта={simple_result[1]}")

        # Проверим контракт
        contract_query = "SELECT id, id_customer, date_contract, number_contract, theme_contract FROM gen_report_context_contracts.contracts WHERE id = %s"
        cursor.execute(contract_query, (simple_result[1],))
        contract_result = cursor.fetchone()

        if contract_result:
            print(f"Контракт найден: ID={contract_result[0]}, ID заказчика={contract_result[1]}")
            print(f"Дата контракта: {contract_result[2]}, Номер: {contract_result[3]}, Тема: {contract_result[4]}")

            # Получим данные заявки
            request_query = "SELECT application_number FROM gen_report_context_contracts.requests WHERE id = %s"
            cursor.execute(request_query, (simple_result[2],))  # simple_result[2] это id_requests
            request_result = cursor.fetchone()
            app_number = request_result[0] if request_result else "Номер не найден"
            print(f"Номер заявки: {app_number}")

            # Проверим организацию
            org_query = "SELECT id, position_dative, long_name_organisation, representative_name_short_dative, representative_appeal FROM gen_report_context_contracts.organizations WHERE id = %s"
            cursor.execute(org_query, (contract_result[1],))
            org_result = cursor.fetchone()

            if org_result:
                print(f"Организация найдена: ID={org_result[0]}, position_dative={org_result[1]}")
                print(
                    f"long_name_organisation={org_result[2]}, representative_name_short_dative={org_result[3]}, representative_appeal={org_result[4]}")

                # Получаем тексты из textforformdocument
                text_queries = [
                    ("soprovod_1", "text_soprovod_1"),
                    ("soprovod_3", "text_soprovod_3"),
                    ("soprovod_4", "text_soprovod_4"),
                    ("soprovod_5", "text_soprovod_5")
                ]

                texts = {}
                for key, var_name in text_queries:
                    text_query = f"SELECT text_data FROM gen_report_context_contracts.textforformdocument WHERE key = '{key}'"
                    cursor.execute(text_query)
                    text_result = cursor.fetchone()
                    texts[var_name] = text_result[0] if text_result else f"Текст {key} не найден"
                    print(f"Текст {key}: {texts[var_name]}")

                # Получаем данные об исполнителе (id_contractor)
                contractor_query = "SELECT id_contractor FROM gen_report_context_contracts.contracts WHERE id = %s"
                cursor.execute(contractor_query, (contract_result[0],))
                contractor_id_result = cursor.fetchone()

                contractor_data = {}
                if contractor_id_result:
                    contractor_id = contractor_id_result[0]
                    contractor_org_query = "SELECT position_nominative, long_name_organisation, representative_signature FROM gen_report_context_contracts.organizations WHERE id = %s"
                    cursor.execute(contractor_org_query, (contractor_id,))
                    contractor_org_result = cursor.fetchone()

                    if contractor_org_result:
                        contractor_data = {
                            'contractor_position_nominative': contractor_org_result[0],
                            'contractor_long_name_organisation': contractor_org_result[1],
                            'contractor_representative_signature': contractor_org_result[2]
                        }
                        print(f"Исполнитель найден: {contractor_data}")
                    else:
                        print(f"Организация исполнителя с ID {contractor_id} не найдена")
                else:
                    print("ID исполнителя не найден")

                return {
                    'report_id': simple_result[0],
                    'contract_id': contract_result[0],
                    'customer_id': contract_result[1],
                    'position_dative': org_result[1],
                    'long_name_organisation': org_result[2],
                    'representative_name_short_dative': org_result[3],
                    'representative_appeal': org_result[4],
                    'date_contract': contract_result[2],
                    'number_contract': contract_result[3],
                    'theme_contract': contract_result[4],
                    'application_number': app_number,
                    **texts,
                    **contractor_data
                }
            else:
                print(f"Организация с ID {contract_result[1]} не найдена")
        else:
            print(f"Контракт с ID {simple_result[1]} не найден")

        return None

    except Exception as e:
        print(f"Ошибка при получении данных: {e}")
        return None
    finally:
        cursor.close()


def create_word_document(data, report_id):
    """Создание Word документа с форматированием"""
    try:
        print(f"Создаем документ с данными: {data}")

        doc = Document()

        # 1. Добавляем основной текст с отступом слева (сдвигаем вправо)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Устанавливаем отступ слева (сдвигаем текст вправо на 4 см)
        left_indent = Inches(4)
        paragraph.paragraph_format.left_indent = left_indent

        # Добавляем основной текст с заглавной буквы
        main_text = f"{data['position_dative']} {data['long_name_organisation']} {data['representative_name_short_dative']}"
        # Делаем первую букву заглавной
        if main_text:
            main_text = main_text[0].upper() + main_text[1:]
        run = paragraph.add_run(main_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        print(f"Добавлен основной текст: '{main_text}'")

        # 2. Добавляем дату (без отступа)
        date_paragraph = doc.add_paragraph()
        current_date = datetime.now().strftime("%d.%m.%Y")
        date_run = date_paragraph.add_run(current_date)
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        print(f"Добавлена дата: '{current_date}'")

        # 3. Добавляем текст из textforformdocument (по левому краю)
        text_paragraph = doc.add_paragraph()
        text_run = text_paragraph.add_run(data['text_soprovod_5'])
        text_run.font.name = 'Times New Roman'
        text_run.font.size = Pt(12)
        print(f"Добавлен текст soprovod_5: '{data['text_soprovod_5']}'")

        # 4. Добавляем representative_appeal (по центру)
        appeal_paragraph = doc.add_paragraph()
        appeal_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        appeal_run = appeal_paragraph.add_run(data['representative_appeal'])
        appeal_run.font.name = 'Times New Roman'
        appeal_run.font.size = Pt(12)
        print(f"Добавлен representative_appeal: '{data['representative_appeal']}'")

        # 5. Добавляем большое предложение с абзацным отступом
        big_paragraph = doc.add_paragraph()
        big_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине

        # Устанавливаем абзацный отступ (как таб)
        big_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        # Формируем дату в нужном формате
        if data['date_contract']:
            contract_date = data['date_contract']
            if isinstance(contract_date, str):
                contract_date = datetime.strptime(contract_date, '%Y-%m-%d').date()
            day = contract_date.day
            month_name = get_month_name(contract_date.month)
            year = contract_date.year
            date_str = f"от {day} «{month_name}» {year} г."
        else:
            date_str = "дата не указана"

        # Формируем большое предложение
        big_text = f"{data['text_soprovod_1']} {date_str} № {data['number_contract']} {data['theme_contract']} {data['text_soprovod_3']} № {data['application_number']}:"

        big_run = big_paragraph.add_run(big_text)
        big_run.font.name = 'Times New Roman'
        big_run.font.size = Pt(12)
        print(f"Добавлено большое предложение: '{big_text}'")

        # 6. Добавляем текст soprovod_4
        soprovod4_paragraph = doc.add_paragraph()
        soprovod4_run = soprovod4_paragraph.add_run(data['text_soprovod_4'])
        soprovod4_run.font.name = 'Times New Roman'
        soprovod4_run.font.size = Pt(12)
        print(f"Добавлен текст soprovod_4: '{data['text_soprovod_4']}'")

        # 7. Добавляем 8 пустых строк
        for i in range(8):
            doc.add_paragraph()

        # 8. Добавляем подпись исполнителя с выравниванием
        signature_paragraph = doc.add_paragraph()
        signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по левому краю

        # Формируем текст подписи
        left_part = f"{data['contractor_position_nominative']} {data['contractor_long_name_organisation']}"
        right_part = data['contractor_representative_signature']

        # Добавляем левую часть
        left_run = signature_paragraph.add_run(left_part)
        left_run.font.name = 'Times New Roman'
        left_run.font.size = Pt(12)

        # Динамически рассчитываем количество пробелов с учетом шрифта
        # Увеличиваем расчеты в 2 раза для правильного выравнивания
        char_width = 14.4  # пикселей на символ (увеличено в 2 раза)
        page_width = 1190  # пикселей ширина страницы A4 (увеличено в 2 раза)
        available_width = page_width - 200  # отступы слева и справа (увеличено в 2 раза)

        left_width = len(left_part) * char_width
        right_width = len(right_part) * char_width
        spaces_width = available_width - left_width - right_width

        # Конвертируем ширину в количество пробелов
        space_width = char_width  # пробел имеет такую же ширину как обычный символ
        spaces_needed = int(spaces_width / space_width)

        # Увеличиваем количество пробелов в 2 раза
        spaces_needed = spaces_needed * 2

        # Минимум 20 пробелов, максимум 100
        spaces_count = max(20, min(100, spaces_needed))
        spaces = " " * spaces_count

        signature_paragraph.add_run(spaces)
        print(
            f"Левая часть: {len(left_part)} символов ({left_width:.1f}px), правая: {len(right_part)} символов ({right_width:.1f}px)")
        print(f"Доступная ширина: {available_width}px, пробелов нужно: {spaces_needed}, используется: {spaces_count}")

        # Добавляем правую часть
        right_run = signature_paragraph.add_run(right_part)
        right_run.font.name = 'Times New Roman'
        right_run.font.size = Pt(12)

        print(f"Добавлена подпись: '{left_part}' -> '{right_part}'")

        # Сохраняем документ с уникальным именем
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # output_path = f"Результат/сопровод_отчет_{REPORT_ID}_{timestamp}.docx"
        output_path = f"{report_id}/сопровод_{timestamp}.docx"

        file = io.BytesIO()
        doc.save(file)
        file.seek(0)
        print(f"Документ создан.")
        return file, output_path

    except Exception as e:
        print(f"Ошибка при создании документа: {e}")


def generate_soprovod(report_id):
    """Основная функция"""
    print(f"Обработка отчета ID: {report_id}")

    # Подключение к БД
    conn = connect_to_db()
    if not conn:
        return

    try:
        # Получение данных
        data = get_report_data(conn, report_id)
        if not data:
            return

        print(f"Найден отчет: {data['report_id']}")
        print(f"ID контракта: {data['contract_id']}")
        print(f"ID заказчика: {data['customer_id']}")
        print(f"Должность (дательный падеж): {data['position_dative']}")

        # Создание Word документа
        if data['position_dative']:
            # создание файла
            file, file_name = create_word_document(data, report_id)
        else:
            print("Поле position_dative пустое")

    except Exception as e:
        print(f"Общая ошибка: {e}")
        raise e
    finally:
        conn.close()
    return file, file_name


if __name__ == "__main__":
    minio_client = Minio(
        endpoint=os.getenv('S3_ENDPOINT_URL', 'minio.upk-mos.ru'),
        access_key=os.getenv('S3_ACCESS_KEY'),
        secret_key=os.getenv('S3_SECRET_KEY'),
        secure=os.getenv('S3_SECURE', 'False').lower() == 'true'
    )
    print(generate_soprovod(16, minio_client))
