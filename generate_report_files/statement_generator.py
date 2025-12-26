import io
import os
import psycopg2
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from dotenv import load_dotenv

from utils.postprocessing_report_file import upload_to_s3, write_s3path_to_bd

# Загружаем переменные из .env файла
load_dotenv()

# ID отчета для обработки (задавайте вручную)
# REPORT_ID = 15

# Параметры подключения к БД
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'database': os.getenv('DB_NAME', 'your_database'),
    'user': os.getenv('DB_USER', 'your_user'),
    'password': os.getenv('DB_PASSWORD', 'your_password'),
    'port': os.getenv('DB_PORT', '5432')
}

def connect_to_db():
    """Подключение к базе данных с обработкой ошибок"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"Ошибка подключения к БД: {e}")
        return None

def get_vedomost_data(conn, report_id):
    """Получение данных для ведомости"""
    try:
        cursor = conn.cursor()
        
        # Сначала проверим, есть ли отчет вообще
        check_query = "SELECT id FROM gen_report_context_contracts.reports WHERE id = %s"
        cursor.execute(check_query, (report_id,))
        report_exists = cursor.fetchone()
        
        if not report_exists:
            print(f"Отчет с ID {report_id} не найден в таблице reports")
            return None
            
        print(f"Отчет {report_id} найден, проверяем связи...")
        
        # Получим данные отчета
        simple_query = "SELECT id, id_contracts FROM gen_report_context_contracts.reports WHERE id = %s"
        cursor.execute(simple_query, (report_id,))
        simple_result = cursor.fetchone()
        
        if not simple_result:
            print("Не удалось получить данные отчета")
            return None
            
        print(f"Данные отчета: ID={simple_result[0]}, ID контракта={simple_result[1]}")
        
        # Проверим контракт
        contract_query = "SELECT id, id_customer, id_contractor, date_contract, number_contract, theme_contract FROM gen_report_context_contracts.contracts WHERE id = %s"
        cursor.execute(contract_query, (simple_result[1],))
        contract_result = cursor.fetchone()
        
        if contract_result:
            print(f"Контракт найден: ID={contract_result[0]}, ID заказчика={contract_result[1]}, ID исполнителя={contract_result[2]}")
            print(f"Дата контракта: {contract_result[3]}, Номер: {contract_result[4]}, Тема: {contract_result[5]}")
            
            # Получаем тексты из textforformdocument по ключам vedomost_1-vedomost_11
            text_queries = [
                ("vedomost_1", "vedomost_1_text"),
                ("vedomost_2", "vedomost_2_text"),
                ("vedomost_3", "vedomost_3_text"),
                ("vedomost_4", "vedomost_4_text"),
                ("vedomost_5", "vedomost_5_text"),
                ("vedomost_6", "vedomost_6_text"),
                ("vedomost_7", "vedomost_7_text"),
                ("vedomost_8", "vedomost_8_text"),
                ("vedomost_9", "vedomost_9_text"),
                ("vedomost_10", "vedomost_10_text"),
                ("vedomost_11", "vedomost_11_text")
            ]
            
            texts = {}
            for key, var_name in text_queries:
                text_query = f"SELECT text_data FROM gen_report_context_contracts.textforformdocument WHERE key = '{key}'"
                cursor.execute(text_query)
                text_result = cursor.fetchone()
                texts[var_name] = text_result[0] if text_result else f"Текст {key} не найден"
                print(f"Текст {key}: {texts[var_name]}")
            
            # Получаем данные о заказчике (id_customer)
            customer_query = "SELECT long_name_organisation, position_nominative, representative_signature FROM gen_report_context_contracts.organizations WHERE id = %s"
            cursor.execute(customer_query, (contract_result[1],))
            customer_result = cursor.fetchone()
            
            if customer_result:
                print(f"Заказчик найден: {customer_result[0]}")
            else:
                print(f"Заказчик с ID {contract_result[1]} не найден")
                customer_result = (None, None, None)
            
            # Получаем данные об исполнителе (id_contractor)
            contractor_query = "SELECT long_name_organisation, position_nominative, representative_signature FROM gen_report_context_contracts.organizations WHERE id = %s"
            cursor.execute(contractor_query, (contract_result[2],))
            contractor_result = cursor.fetchone()
            
            if contractor_result:
                print(f"Исполнитель найден: {contractor_result[0]}")
            else:
                print(f"Исполнитель с ID {contract_result[2]} не найден")
                contractor_result = (None, None, None)
            
            # Получаем данные из таблицы requests для ведомости
            requests_query = """
                SELECT 
                    media_carrier_type_id,
                    media_material_name_1,
                    media_file_name_1,
                    media_material_name_2,
                    media_file_name_2
                FROM gen_report_context_contracts.requests 
                WHERE id_contracts = %s
            """
            cursor.execute(requests_query, (contract_result[0],))
            requests_result = cursor.fetchone()
            
            if requests_result:
                print(f"Данные заявки найдены для ведомости")
                print(f"Тип носителя: {requests_result[0]}")
                print(f"Материал 1: {requests_result[1]}, Файл 1: {requests_result[2]}")
                print(f"Материал 2: {requests_result[3]}, Файл 2: {requests_result[4]}")
            else:
                print(f"Данные заявки для контракта {contract_result[0]} не найдены")
                requests_result = (None, None, None, None, None)
            
            return {
                'report_id': simple_result[0],
                'contract_id': contract_result[0], 
                'customer_id': contract_result[1],
                'contractor_id': contract_result[2],
                'date_contract': contract_result[3],
                'number_contract': contract_result[4],
                'theme_contract': contract_result[5],
                'customer_long_name_organisation': customer_result[0],
                'customer_position_nominative': customer_result[1],
                'customer_representative_signature': customer_result[2],
                'contractor_long_name_organisation': contractor_result[0],
                'contractor_position_nominative': contractor_result[1],
                'contractor_representative_signature': contractor_result[2],
                'media_carrier_type_id': requests_result[0],
                'media_material_name_1': requests_result[1],
                'media_file_name_1': requests_result[2],
                'media_material_name_2': requests_result[3],
                'media_file_name_2': requests_result[4],
                **texts
            }
        else:
            print(f"Контракт с ID {simple_result[1]} не найден")
            
        return None
            
    except Exception as e:
        print(f"Ошибка при получении данных: {e}")
        return None
    finally:
        cursor.close()

def create_vedomost_document(data, report_id):
    """Создание ведомости в формате Word"""
    try:
        print(f"Создаем ведомость с данными: {data}")
        
        doc = Document()
        
        # 1. Добавляем заголовок с long_name_organisation по центру курсивом
        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if data['customer_long_name_organisation']:
            header_run = header_paragraph.add_run(data['customer_long_name_organisation'])
            header_run.font.name = 'Times New Roman'
            header_run.font.size = Pt(12)
            header_run.italic = True  # Курсив
            print(f"Добавлен заголовок: '{data['customer_long_name_organisation']}' (курсив)")
        else:
            header_run = header_paragraph.add_run("Наименование организации не указано")
            header_run.font.name = 'Times New Roman'
            header_run.font.size = Pt(12)
            header_run.italic = True
        
        # 2. Добавляем пустую строку после заголовка
        doc.add_paragraph()
        
        # 3. Создаем таблицу 2x2
        table = doc.add_table(rows=2, cols=2)
        
        # Убираем видимые границы таблицы
        for row in table.rows:
            for cell in row.cells:
                # Убираем все границы
                cell._tc.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._tc.get_or_add_tcPr().find(qn('w:tcBorders'))
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    borders.append(border)
                
                # Настраиваем выравнивание по центру и середине
                cell.vertical_alignment = 1  # По середине
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing = 1.5
                    # Устанавливаем шрифт для текста в ячейках
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
        
        # Устанавливаем ширину таблицы на всю страницу
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3.25)  # Половина ширины страницы A4
        
        # Заполняем ячейки таблицы
        # Левая верхняя ячейка (customer)
        cell_00 = table.cell(0, 0)
        cell_00.vertical_alignment = 0  # Выравнивание по верху ячейки
        cell_00_paragraph = cell_00.paragraphs[0]
        cell_00_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_00_paragraph.paragraph_format.line_spacing = 1.5
        
        # vedomost_1 (жирным)
        if data['vedomost_1_text']:
            run = cell_00_paragraph.add_run(data['vedomost_1_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cell_00_paragraph.add_run("\n")
        
        # position_nominative
        if data['customer_position_nominative']:
            run = cell_00_paragraph.add_run(data['customer_position_nominative'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            cell_00_paragraph.add_run("\n")
        
        # long_name_organisation
        if data['customer_long_name_organisation']:
            run = cell_00_paragraph.add_run(data['customer_long_name_organisation'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Добавляем 2 пустых строки
        for i in range(2):
            cell_00_paragraph.add_run("\n")
        
        # Правая верхняя ячейка (contractor)
        cell_01 = table.cell(0, 1)
        cell_01.vertical_alignment = 0  # Выравнивание по верху ячейки
        cell_01_paragraph = cell_01.paragraphs[0]
        cell_01_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_01_paragraph.paragraph_format.line_spacing = 1.5
        
        # vedomost_1 (жирным)
        if data['vedomost_1_text']:
            run = cell_01_paragraph.add_run(data['vedomost_1_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cell_01_paragraph.add_run("\n")
        
        # position_nominative
        if data['contractor_position_nominative']:
            run = cell_01_paragraph.add_run(data['contractor_position_nominative'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            cell_01_paragraph.add_run("\n")
        
        # long_name_organisation
        if data['contractor_long_name_organisation']:
            run = cell_01_paragraph.add_run(data['contractor_long_name_organisation'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Добавляем 2 пустых строки
        for i in range(2):
            cell_01_paragraph.add_run("\n")
        
        # Левая нижняя ячейка (customer) с подписями
        cell_10 = table.cell(1, 0)
        cell_10_paragraph = cell_10.paragraphs[0]
        cell_10_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_10_paragraph.paragraph_format.line_spacing = 1.5
        
        # Подпись заказчика: "___________________/ {representative_signature}"
        if data['customer_representative_signature']:
            run = cell_10_paragraph.add_run("___________________/ ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run = cell_10_paragraph.add_run(data['customer_representative_signature'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            run = cell_10_paragraph.add_run("___________________/ ________________")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        cell_10_paragraph.add_run("\n")
        
        # Дата: "« » ____________  2025 г."
        run = cell_10_paragraph.add_run("« » ____________  2025 г.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cell_10_paragraph.add_run("\n")
        
        # vedomost_2
        if data['vedomost_2_text']:
            run = cell_10_paragraph.add_run(data['vedomost_2_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Правая нижняя ячейка (contractor) с подписями
        cell_11 = table.cell(1, 1)
        cell_11_paragraph = cell_11.paragraphs[0]
        cell_11_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_11_paragraph.paragraph_format.line_spacing = 1.5
        
        # Подпись исполнителя: "___________________/{representative_signature}"
        if data['contractor_representative_signature']:
            run = cell_11_paragraph.add_run("___________________/")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run = cell_11_paragraph.add_run(data['contractor_representative_signature'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            run = cell_11_paragraph.add_run("___________________/________________")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        cell_11_paragraph.add_run("\n")
        
        # Текущая дата: "«4» июля 2025 г."
        current_date = datetime.now()
        day = current_date.day
        month_name = get_month_name(current_date.month)
        year = current_date.year
        date_text = f"«{day}» {month_name} {year} г."
        run = cell_11_paragraph.add_run(date_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cell_11_paragraph.add_run("\n")
        
        # vedomost_2
        if data['vedomost_2_text']:
            run = cell_11_paragraph.add_run(data['vedomost_2_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        print(f"Добавлена таблица 2x2 с данными заказчика и исполнителя")
        
        # 4. Добавляем 2 пустых строки после таблицы
        for i in range(2):
            doc.add_paragraph()
        
        # 5. Добавляем текст vedomost_3 + дату договора + тему договора по центру
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.line_spacing = 1.5
        title_paragraph.paragraph_format.space_after = Pt(0)  # Убираем интервал между абзацами
        
        # vedomost_3
        if data['vedomost_3_text']:
            run = title_paragraph.add_run(data['vedomost_3_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            title_paragraph.add_run(" ")
        
        # Формируем дату договора в нужном формате: "от 28 «декабря» 2024 г. № 500-976"
        if data['date_contract']:
            contract_date = data['date_contract']
            if isinstance(contract_date, str):
                contract_date = datetime.strptime(contract_date, '%Y-%m-%d').date()
            day = contract_date.day
            month_name = get_month_name(contract_date.month)
            year = contract_date.year
            date_text = f"от {day} «{month_name}» {year} г. № {data['number_contract']}"
        else:
            date_text = "от дата не указана"
        
        run = title_paragraph.add_run(date_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # theme_contract с новой строки (без абзацного отступа)
        if data['theme_contract']:
            theme_paragraph = doc.add_paragraph()
            theme_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            theme_paragraph.paragraph_format.line_spacing = 1.5
            theme_paragraph.paragraph_format.first_line_indent = Inches(0)  # Убираем абзацный отступ
            run = theme_paragraph.add_run(data['theme_contract'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        print(f"Добавлен заголовок с vedomost_3, датой договора и темой договора")
        
        # 6. Добавляем 2 пустых строки
        for i in range(2):
            doc.add_paragraph()
        
        # 7. Добавляем vedomost_4 по центру
        if data['vedomost_4_text']:
            vedomost_4_paragraph = doc.add_paragraph()
            vedomost_4_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vedomost_4_paragraph.paragraph_format.line_spacing = 1.5
            run = vedomost_4_paragraph.add_run(data['vedomost_4_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # 8. Добавляем пустую строку
        doc.add_paragraph()
        
        # 9. Добавляем vedomost_5 по центру
        if data['vedomost_5_text']:
            vedomost_5_paragraph = doc.add_paragraph()
            vedomost_5_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vedomost_5_paragraph.paragraph_format.line_spacing = 1.5
            run = vedomost_5_paragraph.add_run(data['vedomost_5_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # 10. Добавляем 4 пустых строки
        for i in range(4):
            doc.add_paragraph()
        
        # 11. Добавляем vedomost_6 по центру
        if data['vedomost_6_text']:
            vedomost_6_paragraph = doc.add_paragraph()
            vedomost_6_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vedomost_6_paragraph.paragraph_format.line_spacing = 1.5
            run = vedomost_6_paragraph.add_run(data['vedomost_6_text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        print(f"Добавлен блок с vedomost_4, vedomost_5 и vedomost_6")
        
        # 12. Добавляем таблицу 4x3 с данными о материалах
        data_table = doc.add_table(rows=3, cols=4)  # 3 строки, 4 столбца (без первого пустого столбца)
        data_table.style = 'Table Grid'  # Видимые границы
        
        # Настраиваем выравнивание ячеек по центру
        for row in data_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1  # По середине
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing = 1.0  # Интервал 1.0 для нижней таблицы
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
        
        # Заполняем первую строку (заголовки)
        headers = [
            data['vedomost_7_text'] if data['vedomost_7_text'] else "vedomost_7",
            data['vedomost_8_text'] if data['vedomost_8_text'] else "vedomost_8", 
            data['vedomost_9_text'] if data['vedomost_9_text'] else "vedomost_9",
            data['vedomost_10_text'] if data['vedomost_10_text'] else "vedomost_10"
        ]
        
        for i, header in enumerate(headers):
            cell = data_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0  # Интервал 1.0 для заголовков нижней таблицы
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        
        # Заполняем вторую строку
        # Столбец 0 (№ п/п) - vedomost_11, будет объединен
        data_table.cell(1, 0).text = data['vedomost_11_text'] or ""
        
        # Столбец 1 (Тип и №/ID носителя) - данные
        data_table.cell(1, 1).text = data['media_carrier_type_id'] or ""
        
        # Столбец 2 (Наименование материала) - материал 1
        data_table.cell(1, 2).text = data['media_material_name_1'] or ""
        
        # Столбец 3 (Имя файла) - файл 1 (для первой строки)
        data_table.cell(1, 3).text = data['media_file_name_1'] or ""
        
        # Заполняем третью строку
        # Столбец 0 (№ п/п) - пустой, будет объединен
        data_table.cell(2, 0).text = ""
        
        # Столбец 1 (Тип и №/ID носителя) - пустой, будет объединен
        data_table.cell(2, 1).text = ""
        
        # Столбец 2 (Наименование материала) - материал 2
        data_table.cell(2, 2).text = data['media_material_name_2'] or ""
        
        # Столбец 3 (Имя файла) - файл 2
        data_table.cell(2, 3).text = data['media_file_name_2'] or ""
        
        # Объединяем ячейки в столбцах № п/п и Тип и №/ID носителя (строки 2-3)
        # Объединяем ячейки (1,0) и (2,0) - столбец № п/п
        cell_1_0 = data_table.cell(1, 0)
        cell_2_0 = data_table.cell(2, 0)
        cell_1_0.merge(cell_2_0)
        
        # Объединяем ячейки (1,1) и (2,1) - столбец Тип и №/ID носителя
        cell_1_1 = data_table.cell(1, 1)
        cell_2_1 = data_table.cell(2, 1)
        cell_1_1.merge(cell_2_1)
        
        # Применяем шрифт ко всем ячейкам таблицы после заполнения
        for row in data_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing = 1.0  # Интервал 1.0 для нижней таблицы
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
        
        print(f"Добавлена таблица с данными о материалах")
        
        # Сохраняем документ с уникальным именем
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # output_path = f"Результат/ведомость_отчет_{report_id}_{timestamp}.docx"
        output_path = f'{report_id}/ведомость_{timestamp}.docx'
        file = io.BytesIO()
        doc.save(file)
        file.seek(0)
        print(f"Ведомость создана.")
        return file, output_path
        
    except Exception as e:
        print(f"Ошибка при создании ведомости: {e}")

def get_month_name(month_num):
    """Преобразование номера месяца в название"""
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    return months.get(month_num, "месяца")

def generate_vedomost(report_id):
    """Основная функция"""
    print(f"Обработка отчета ID: {report_id}")
    
    # Подключение к БД
    conn = connect_to_db()
    if not conn:
        return
    
    try:
        # Получение данных
        data = get_vedomost_data(conn, report_id)
        if not data:
            return
            
        print(f"Найден отчет: {data['report_id']}")
        print(f"ID контракта: {data['contract_id']}")
        print(f"ID заказчика: {data['customer_id']}")
        
        # Создание ведомости
        file, file_name = create_vedomost_document(data, report_id)

        # отправка в S3 и запись пути к файлу в БД
        # s3_filepath = upload_to_s3(minio_client, file, file_name)

        # запись пути к файлу в БД
        # write_s3path_to_bd(report_id, 'file_machine_media_statement', s3_filepath)
            
    except Exception as e:
        print(f"Общая ошибка: {e}")
        raise e
    finally:
        conn.close()

    return file, file_name

if __name__ == "__main__":
    print(generate_vedomost(16))
