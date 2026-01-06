#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор отчетов в формате Word
"""
import io
import os
import sys
import json
import tempfile
import requests
import psycopg2
from psycopg2.extensions import cursor as cur
import time
from minio import Minio
from dotenv import load_dotenv
from typing import Dict, List, Optional, Any
from datetime import datetime
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn


# Загружаем переменные окружения
load_dotenv()

# Настройки для таблицы лучших объявлений
TOP_ADS_COUNT = 2  # Количество лучших объявлений для отображения в таблице

class ReportGenerator:
    def __init__(self):
        """Инициализация подключений к БД и MinIO"""
        # Настройки БД
        self.db_config = {
            'host': os.getenv('DB_HOST', 'localhost'),
            'port': os.getenv('DB_PORT', '5432'),
            'database': os.getenv('DB_NAME'),
            'user': os.getenv('DB_USER'),
            'password': os.getenv('DB_PASSWORD')
        }
        
        # Настройки MinIO
        self.minio_client = Minio(
            endpoint=os.getenv('S3_ENDPOINT_URL', 'minio.upk-mos.ru'),
            access_key=os.getenv('S3_ACCESS_KEY'),
            secret_key=os.getenv('S3_SECRET_KEY'),
            secure=os.getenv('S3_SECURE', 'False').lower() == 'true'
        )
        
        self.bucket_name = os.getenv('S3_BUCKET_NAME', 'dit-services-dev')
        
        # Папка для результатов
        self.output_folder = 'report_results'
        self._ensure_output_folder()

    def _ensure_output_folder(self):
        """Создать папку для результатов, если её нет"""
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
            print(f"📁 Создана папка для результатов: {self.output_folder}")
        else:
            print(f"📁 Папка для результатов: {self.output_folder}")

    def _connect_to_db(self, max_retries: int = 3, initial_delay: float = 1.0):
        """Подключение к БД с повторными попытками
        
        Args:
            max_retries: Максимальное количество попыток подключения
            initial_delay: Начальная задержка между попытками в секундах (будет увеличиваться экспоненциально)
            
        Returns:
            Объект соединения с БД или None при неудаче
        """
        delay = initial_delay
        last_error = None
        
        for attempt in range(1, max_retries + 1):
            try:
                conn = psycopg2.connect(**self.db_config)
                if attempt > 1:
                    print(f"✅ Успешное подключение к БД после {attempt} попыток")
                return conn
            except Exception as e:
                last_error = e
                if attempt < max_retries:
                    print(f"⚠️ Попытка {attempt}/{max_retries} подключения к БД не удалась: {e}")
                    print(f"⏳ Повторная попытка через {delay:.1f} сек...")
                    time.sleep(delay)
                    delay *= 2  # Экспоненциальное увеличение задержки
        
        print(f"❌ Не удалось подключиться к БД после {max_retries} попыток. Последняя ошибка: {last_error}")
        return None

    def get_report_text(self, key: str) -> Optional[str]:
        """Получить текст для отчета по ключу из БД"""
        conn = None
        try:
            conn = self._connect_to_db()
            if not conn:
                return None
                
            cursor: cur = conn.cursor()
            
            # Устанавливаем схему по умолчанию
            cursor.execute("SET search_path TO gen_report_context_contracts, public;")
            
            # Получаем текст по ключу
            query = """
            SELECT text_data
            FROM textforformdocument
            WHERE key = %s
            """
            
            cursor.execute(query, (key,))
            result = cursor.fetchone()
            
            cursor.close()
            conn.close()
            
            return result[0] if result else None
            
        except Exception as e:
            # !!! здесь иногда возникает ошибка: server closed the connection unexpectedly
            #         This probably means the server terminated abnormally
            #         before or while processing the request.
            # пока не удалось понять почему
            print(f"❌ Ошибка при получении текста отчета: {e}")
            if conn:
                cursor.close()
                conn.close()
            # return None
            raise e

    def get_report_data(self, report_id: int) -> Optional[Dict]:
        """Получить данные отчета из БД"""
        conn = None
        try:
            conn = self._connect_to_db()
            if not conn:
                return None
                
            cursor = conn.cursor()
            
            # Устанавливаем схему по умолчанию
            cursor.execute("SET search_path TO gen_report_context_contracts, public;")
            
            # Получаем данные отчета
            query = """
            SELECT 
                r.id,
                r.id_contracts,
                r.id_requests,
                c.number_contract,
                c.date_contract,
                c.theme_contract,
                req.date_request,
                req.application_number,
                c.id_customer,
                c.id_contractor,
                cust.long_name_organisation as customer_org,
                cust.position_nominative as customer_position,
                cust.representative_signature as customer_signature,
                contr.long_name_organisation as contractor_org,
                contr.position_nominative as contractor_position,
                contr.representative_signature as contractor_signature,
                req.digital_project,
                req.start_date,
                req.end_date,
                req.list_recommended_campaign_types,
                req.list_recommended_formats_ads,
                req.description_target_audience,
                req.interests,
                req.examples_published_ads,
                req.conclusions_recommendations,
                c.goals,
                c.tasks,
                c.description_target_audience,
                c.requirements_visual_materials,
                c.requirements_text_materials,
                c.kpi_plan_clicks,
                c.kpi_plan_reject,
                req.campany_yandex_direct
            FROM reports r
            JOIN contracts c ON r.id_contracts = c.id
            JOIN requests req ON r.id_requests = req.id
            JOIN organizations cust ON c.id_customer = cust.id
            JOIN organizations contr ON c.id_contractor = contr.id
            WHERE r.id = %s
            """
            
            cursor.execute(query, (report_id,))
            result = cursor.fetchone()
            
            if not result:
                print(f"❌ Отчет с ID {report_id} не найден")
                return None
            
            # Отладочная информация
            print(f"🔍 Отладочная информация для отчета {report_id}:")
            print(f"   ID отчета: {result[0]}")
            print(f"   ID контракта: {result[1]}")
            print(f"   ID заявки: {result[2]}")
            print(f"   Номер контракта: {result[3]}")
            print(f"   Дата контракта: {result[4]}")
            print(f"   Тема контракта: {result[5]}")
            print(f"   Дата заявки: {result[6]}")
            print(f"   Номер заявки: {result[7]}")  # Это application_number
            print(f"   Digital project: {result[19]}")
            print(f"   List recommended campaign types: {result[20]}")
            print(f"   List recommended formats ads: {result[21]}")
            print(f"   Description target audience (requests): {result[22]}")
            print(f"   Interests: {result[23]}")
                
            # Получаем термины для контракта
            terms_query = """
            SELECT term_title, term_description
            FROM gen_report_context_contracts.terms
            WHERE id_contract = %s AND (is_deleted = false OR is_deleted IS NULL)
            ORDER BY serial_number
            """
            cursor.execute(terms_query, (result[1],))
            terms = cursor.fetchall()

            # Получаем переписку для заявки
            correspondence_query = """
            SELECT 
                wc.id,
                wc.id_letter_name,
                wc.date_sent,
                wc.file_link,
                tl.theme
            FROM gen_report_context_contracts.workСorrespondence wc
            LEFT JOIN gen_report_context_contracts.themesletter tl ON wc.id_letter_name = tl.id
            WHERE wc.id_requests = %s AND (wc.is_deleted = false OR wc.is_deleted IS NULL)
            ORDER BY wc.date_sent
            """
            cursor.execute(correspondence_query, (result[2],))
            correspondence = cursor.fetchall()

            report_data = {
                'id': result[0],
                'id_contracts': result[1],
                'id_requests': result[2],
                'number_contract': result[3],
                'date_contract': result[4],
                'theme_contract': result[5],
                'date_request': result[6],
                'application_number': result[7],  # Номер заявки из таблицы requests
                'id_customer': result[8],
                'id_contractor': result[9],
                'customer_org': result[10],
                'customer_position': result[11],
                'customer_signature': result[12],
                'contractor_org': result[13],
                'contractor_position': result[14],
                'contractor_signature': result[15],
                'digital_project': result[16],
                'start_date': result[17],
                'end_date': result[18],
                'list_recommended_campaign_types': result[19],
                'list_recommended_formats_ads': result[20],
                'description_target_audience_requests': result[21],
                'interests': result[22],
                'examples_published_ads': result[23],
                'conclusions_recommendations': result[24],
                'goals': result[25],
                'tasks': result[26],
                'description_target_audience': result[27],
                'requirements_visual_materials': result[28],
                'requirements_text_materials': result[29],
                'kpi_plan_clicks': result[30],
                'kpi_plan_reject': result[31],
                'campany_yandex_direct': result[32],
                'terms': terms,
                'correspondence': correspondence
            }
            
            cursor.close()
            conn.close()
            
            return report_data
            
        except Exception as e:
            print(f"❌ Ошибка при получении данных отчета: {e}")
            if conn:
                conn.close()
            # return None
            raise e

    def get_month_name(self, month_num: int) -> str:
        """Преобразование номера месяца в название"""
        months = {
            1: "января", 2: "февраля", 3: "марта", 4: "апреля",
            5: "мая", 6: "июня", 7: "июля", 8: "августа",
            9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
        }
        return months.get(month_num, "месяца")

    def format_date(self, date) -> str:
        """Форматировать дату в нужный формат"""
        if not date:
            return ""
        return f"«{date.day}» {self.get_month_name(date.month)} {date.year} г."

    def create_section_1(self, doc: Document, report_data: Dict) -> None:
        """Создать первый раздел отчета"""
        # Получаем тексты из БД
        report_1 = self.get_report_text('report_1')
        report_2 = self.get_report_text('report_2')
        report_3 = self.get_report_text('report_3')
        
        if not all([report_1, report_2, report_3]):
            print("❌ Не удалось получить все необходимые тексты для раздела 1")
            return

        # Добавляем разрыв страницы, если это не первый раздел
        if len(doc.paragraphs) > 0:
            doc.add_page_break()

        # Добавляем пустую строку перед первым текстом
        doc.add_paragraph()

        # 1. Первая строка с датой и номером контракта (разделена на две строки)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"{report_1} {self.format_date(report_data['date_contract'])} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.bold = True

        # Тема контракта на новой строке
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.bold = True

        # 2. Пустая строка
        doc.add_paragraph()

        # 3. report_2 по центру
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_2)
        run.bold = True

        # 4. report_3 с датой и номером заявки
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Изменено на CENTER
        text = f"{report_3} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = p.add_run(text)
        run.bold = True

        # 5. Четыре пустых строки
        for _ in range(4):
            doc.add_paragraph()

    def create_section_2(self, doc: Document, report_data: Dict) -> None:
        """Создать второй раздел отчета с таблицей"""
        # Получаем тексты из БД
        report_4 = self.get_report_text('report_4')
        report_5 = self.get_report_text('report_5')
        report_6 = self.get_report_text('report_6')
        
        if not all([report_4, report_5, report_6]):
            print("❌ Не удалось получить все необходимые тексты для раздела 2")
            return

        # Создаем таблицу 2x2
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Normal Table'  # Используем стиль без границ
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы на всю страницу
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3.5)  # Половина ширины страницы A4

        # 1. Левая верхняя ячейка (customer)
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # report_4
        run = p.add_run(report_4)
        run.bold = True
        p.add_run('\n')
        
        # long_name_organisation
        run = p.add_run(report_data['customer_org'])
        run.bold = True
        p.add_run('\n')
        
        # position_nominative
        run = p.add_run(report_data['customer_position'])
        run.bold = True
        # Добавляем две пустые строки
        p.add_run('\n\n')

        # 2. Правая верхняя ячейка (contractor)
        cell = table.cell(0, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # report_5
        run = p.add_run(report_5)
        run.bold = True
        p.add_run('\n')
        
        # long_name_organisation
        run = p.add_run(report_data['contractor_org'])
        run.bold = True
        p.add_run('\n')
        
        # position_nominative
        run = p.add_run(report_data['contractor_position'])
        run.bold = True
        # Добавляем две пустые строки
        p.add_run('\n\n')

        # 3. Левая нижняя ячейка (customer)
        cell = table.cell(1, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Подпись
        p.add_run('___________________/ ')
        run = p.add_run(report_data['customer_signature'])
        run.bold = True
        p.add_run('/\n')
        
        # report_6
        p.add_run(report_6)

        # 4. Правая нижняя ячейка (contractor)
        cell = table.cell(1, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Подпись
        p.add_run('___________________/ ')
        run = p.add_run(report_data['contractor_signature'])
        run.bold = True
        p.add_run('/\n')
        
        # report_6
        p.add_run(report_6)

    def create_section_3(self, doc: Document) -> None:
        """Создать второй раздел отчета"""
        # Получаем текст из БД
        report_7 = self.get_report_text('report_7')
        
        if not report_7:
            print("❌ Не удалось получить текст для раздела 2")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # Добавляем заголовок
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']  # Устанавливаем стиль заголовка 1 уровня
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(report_7)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет

        # Добавляем пустую строку перед оглавлением
        doc.add_paragraph()

        # Добавляем автособираемое оглавление
        paragraph = doc.add_paragraph()
        
        # Создаем правильное поле TOC для Word
        run = paragraph.add_run()
        
        # Начало поля
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(begin)
        
        # Инструкция поля - используем стандартный TOC
        instr = OxmlElement('w:instrText')
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instr)
        
        # Разделитель поля
        separate = OxmlElement('w:fldChar')
        separate.set(qn('w:fldCharType'), 'separate')
        run._r.append(separate)
        
        # Текст по умолчанию (будет заменен при обновлении)
        default_text = OxmlElement('w:t')
        default_text.text = "Щелкните правой кнопкой мыши и выберите \"Обновить поле\" для обновления оглавления."
        run._r.append(default_text)
        
        # Конец поля
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        run._r.append(end)

        # Добавляем пустую строку после оглавления
        doc.add_paragraph()

    def create_section_4(self, doc: Document, report_data: Dict) -> None:
        """Создать третий раздел отчета с таблицей терминов"""
        # Получаем тексты из БД
        report_8 = self.get_report_text('report_8')
        report_9 = self.get_report_text('report_9')
        report_10 = self.get_report_text('report_10')
        
        if not all([report_8, report_9, report_10]):
            print("❌ Не удалось получить все необходимые тексты для раздела 3")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # Добавляем заголовок
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(report_8)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет

        # Создаем таблицу
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  # Стиль с видимыми границами
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы
        table.width = Inches(6.0)  # Примерная ширина страницы A4
        # Первая колонка 30%, вторая 70%
        table.columns[0].width = Inches(1.8)  # 30% от 6 inches
        table.columns[1].width = Inches(4.2)  # 70% от 6 inches

        # Заголовки таблицы
        header_cells = table.rows[0].cells
        
        # Первая колонка (report_9)
        header_cells[0].text = report_9
        header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        header_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        # Заливка серым цветом #666666
        header_cells[0]._tc.get_or_add_tcPr().append(parse_xml(
            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="666666"/>'))
        
        # Вторая колонка (report_10)
        header_cells[1].text = report_10
        header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        header_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        # Заливка серым цветом #666666
        header_cells[1]._tc.get_or_add_tcPr().append(parse_xml(
            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="666666"/>'))

        # Добавляем данные
        for term_title, term_description in report_data['terms']:
            row = table.add_row()
            cells = row.cells
            
            # Термин
            cells[0].text = term_title
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            
            # Описание
            cells[1].text = term_description
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[1].paragraphs[0].runs[0].font.size = Pt(12)

    def create_section_5(self, doc: Document, report_data: Dict) -> None:
        """Создать четвертый раздел отчета"""
        # Получаем тексты из БД
        report_11 = self.get_report_text('report_11')
        report_21 = self.get_report_text('report_21')
        report_12 = self.get_report_text('report_12')
        
        if not all([report_11, report_21, report_12]):
            print("❌ Не удалось получить все необходимые тексты для раздела 4")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"{report_11} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_21 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_21} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 4. report_12 по левому краю
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_12)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def format_number_with_spaces(self, number):
        """Форматировать число с пробелами между разрядами и запятой для десятичных"""
        if number is None:
            return "0"
        
        # Обрабатываем как float для корректной работы с десятичными
        try:
            num = float(number)
        except (ValueError, TypeError):
            return "0"
        
        # Разделяем целую и дробную части
        if num == int(num):
            # Целое число
            integer_part = int(num)
            decimal_part = None
        else:
            # Число с десятичной частью
            integer_part = int(num)
            decimal_part = num - integer_part
        
        # Форматируем целую часть с пробелами между разрядами
        integer_str = f"{integer_part:,}".replace(",", " ")
        
        # Добавляем десятичную часть с запятой
        if decimal_part is not None:
            # Округляем до 2 знаков после запятой и убираем ведущий 0
            decimal_str = f"{decimal_part:.2f}"[2:]  # Убираем "0."
            # Убираем лишние нули в конце
            decimal_str = decimal_str.rstrip('0')
            if decimal_str:  # Если есть значащие цифры после запятой
                return f"{integer_str},{decimal_str}"
        
        return integer_str

    def format_number_with_two_decimals(self, number):
        """Форматировать число с пробелами между разрядами и всегда 2 знака после запятой"""
        if number is None:
            return "0,00"
        
        # Обрабатываем как float для корректной работы с десятичными
        try:
            num = float(number)
        except (ValueError, TypeError):
            return "0,00"
        
        # Разделяем целую и дробную части
        integer_part = int(num)
        
        # Форматируем целую часть с пробелами между разрядами
        integer_str = f"{integer_part:,}".replace(",", " ")
        
        # Форматируем дробную часть - всегда 2 знака после запятой
        decimal_str = f"{num:.2f}".split('.')[1]  # Берем только дробную часть
        
        return f"{integer_str},{decimal_str}"

    def format_percentage(self, percentage):
        """Форматировать процент без десятичных знаков"""
        if percentage is None:
            return "0"
        return f"{int(percentage)}"

    def load_image_from_minio(self, file_path: str, silent: bool = False) -> bytes:
        """Загрузить изображение из MinIO
        
        Args:
            file_path: Путь к файлу в MinIO
            silent: Если True, не выводить ошибку при отсутствии файла
        """
        try:
            # Обрабатываем случай, когда file_path содержит имя bucket'а
            if file_path.startswith(f"{self.bucket_name}/"):
                # Убираем имя bucket'а из пути
                object_name = file_path[len(f"{self.bucket_name}/"):]
            else:
                # Путь уже содержит только object_name
                object_name = file_path
            
            # Загружаем файл из MinIO
            response = self.minio_client.get_object(
                bucket_name=self.bucket_name,
                object_name=object_name
            )
            
            # Читаем данные
            image_data = response.read()
            response.close()
            response.release_conn()
            
            return image_data
        except Exception as e:
            # Если файл не найден и режим silent, просто возвращаем None
            if silent and 'NoSuchKey' in str(e):
                return None
            print(f"❌ Ошибка загрузки изображения {file_path}: {str(e)}")
            # return None
            raise e
    def check_file_exists_in_minio(self, file_path: str) -> bool:
        """Проверить существование файла в MinIO
        
        Args:
            file_path: Путь к файлу в MinIO
            
        Returns:
            True если файл существует, False если нет
        """
        try:
            # Обрабатываем случай, когда file_path содержит имя bucket'а
            if file_path.startswith(f"{self.bucket_name}/"):
                # Убираем имя bucket'а из пути
                object_name = file_path[len(f"{self.bucket_name}/"):]
            else:
                # Путь уже содержит только object_name
                object_name = file_path
            
            # Проверяем существование объекта
            self.minio_client.stat_object(self.bucket_name, object_name)
            return True
        except Exception:
            return False

    def load_file_from_minio(self, report_id: int, filename: str) -> Optional[Dict]:
        """Загрузить JSON файл из MinIO для конкретного отчета"""
        try:
            folder_path = f"gen_report_context_contracts/data_yandex_direct/{report_id}_результаты"
            object_path = f"{folder_path}/{filename}"
            
            # Проверяем существование объекта
            if self.minio_client.stat_object(self.bucket_name, object_path):
                # Загружаем объект
                response = self.minio_client.get_object(self.bucket_name, object_path)
                content = response.read().decode('utf-8')
                data = json.loads(content)
                print(f"✓ Загружен файл {filename}")
                return data
            else:
                print(f"⚠ Файл {filename} не найден")
                return None
                
        except Exception as e:
            print(f"✗ Ошибка при загрузке {filename}: {e}")
            # return None
            raise e

    def get_unique_images_for_report(self, report_id: int) -> List[Dict]:
        """Получить уникальные изображения для отчета"""
        try:
            # Загружаем данные об изображениях
            images_data = self.load_file_from_minio(report_id, f"image_hashes_report_{report_id}.json")
            if not images_data:
                print("❌ Не удалось загрузить данные об изображениях")
                return []
            
            # Получаем список изображений
            ad_images = images_data.get('result', {}).get('AdImages', [])
            
            # Фильтруем только связанные изображения
            unique_images = []
            seen_hashes = set()
            
            for image in ad_images:
                if (image.get('Associated') == 'YES' and 
                    image.get('AdImageHash') not in seen_hashes):
                    unique_images.append({
                        'hash': image.get('AdImageHash'),
                        'name': image.get('Name', ''),
                        'preview_url': image.get('PreviewUrl', ''),
                        'original_url': image.get('OriginalUrl', ''),
                        'type': image.get('Type', '')
                    })
                    seen_hashes.add(image.get('AdImageHash'))
            
            print(f"✅ Найдено уникальных изображений: {len(unique_images)}")
            return unique_images
            
        except Exception as e:
            print(f"❌ Ошибка при получении изображений: {e}")
            raise e
            # return []

    def get_campaign_project_mapping(self, campany_yandex_direct_json) -> Dict[int, int]:
        """Получить mapping campaign_id -> project_id из поля requests.campany_yandex_direct
        
        Args:
            campany_yandex_direct_json: JSON данные из поля requests.campany_yandex_direct
            
        Returns:
            Словарь {campaign_id: project_id}
        """
        mapping = {}
        try:
            if campany_yandex_direct_json and 'campaigns' in campany_yandex_direct_json:
                for campaign in campany_yandex_direct_json['campaigns']:
                    campaign_id = campaign.get('id')
                    project_id = campaign.get('project_id', 0)
                    if campaign_id:
                        mapping[campaign_id] = project_id
        except Exception as e:
            print(f"⚠️ Ошибка при получении mapping project_id: {e}")
        
        return mapping

    def get_campaigns_data_for_report(self, report_id: int, project_mapping: Dict[int, int] = None) -> List[Dict]:
        """Получить данные о кампаниях для отчета
        
        Args:
            report_id: ID отчета
            project_mapping: Словарь {campaign_id: project_id} для группировки кампаний по проектам
        """
        if project_mapping is None:
            project_mapping = {}
            
        try:
            # Загружаем данные о кампаниях (основной источник порядка)
            campaigns_data = self.load_file_from_minio(report_id, "campaigns.json")
            if not campaigns_data:
                print("❌ Не удалось загрузить данные о кампаниях")
                return []
            
            # Загружаем данные об объявлениях
            ads_data = self.load_file_from_minio(report_id, f"ads_report_{report_id}.json")
            if not ads_data:
                print("❌ Не удалось загрузить данные об объявлениях")
                return []
            
            # Загружаем данные о расширениях
            extensions_data = self.load_file_from_minio(report_id, f"extensions_{report_id}.json")
            if not extensions_data:
                print("❌ Не удалось загрузить данные о расширениях")
                return []
            
            # Загружаем данные о быстрых ссылках
            sitelinks_data = self.load_file_from_minio(report_id, f"sitelinks_{report_id}.json")
            if not sitelinks_data:
                print("❌ Не удалось загрузить данные о быстрых ссылках")
                return []
            
            # Загружаем данные о ключевых фразах
            keywords_data = self.load_file_from_minio(report_id, f"keywords_traffic_forecast_{report_id}.json")
            if not keywords_data:
                print("❌ Не удалось загрузить данные о ключевых фразах")
                return []
            
            # Получаем список объявлений
            ads = ads_data.get('result', {}).get('Ads', [])
            
            # Создаем словарь для связи AdExtensionId с текстом уточнения
            extensions_dict = {}
            # Обрабатываем структуру extensions_1.json (может быть в batch_1 или прямо в result)
            extensions = None
            if 'batch_1' in extensions_data and 'result' in extensions_data['batch_1']:
                extensions = extensions_data['batch_1']['result'].get('AdExtensions', [])
            elif 'result' in extensions_data:
                extensions = extensions_data['result'].get('AdExtensions', [])
            
            if extensions:
                for ext in extensions:
                    if ext.get('Type') == 'CALLOUT' and ext.get('Callout'):
                        ext_id = ext.get('Id')
                        callout_text = ext['Callout'].get('CalloutText')
                        if ext_id and callout_text:
                            extensions_dict[ext_id] = callout_text
            
            # Создаем словарь для связи SitelinkSetId с быстрыми ссылками
            sitelinks_dict = {}
            # Обрабатываем структуру sitelinks файла
            for sitelink_set_id, sitelink_set_data in sitelinks_data.items():
                if isinstance(sitelink_set_data, dict) and 'result' in sitelink_set_data:
                    sitelinks_sets = sitelink_set_data['result'].get('SitelinksSets', [])
                    for sitelinks_set in sitelinks_sets:
                        set_id = sitelinks_set.get('Id')
                        sitelinks = sitelinks_set.get('Sitelinks', [])
                        if set_id and sitelinks:
                            sitelinks_dict[set_id] = sitelinks
            
            # Создаем множество кампаний с ключевыми фразами
            campaigns_with_keywords = set()
            keywords = keywords_data.get('result', {}).get('Keywords', [])
            for keyword in keywords:
                campaign_id = keyword.get('CampaignId')
                keyword_text = keyword.get('Keyword', '')
                
                # Исключаем ключевые фразы с текстом "---autotargeting"
                if campaign_id and keyword_text != "---autotargeting":
                    campaigns_with_keywords.add(campaign_id)
            
            # Получаем список кампаний из campaigns.json в правильном порядке
            campaigns_list_from_json = campaigns_data.get('result', {}).get('Campaigns', [])
            
            # Создаем словарь для быстрого доступа к данным кампаний
            campaigns_dict = {}
            
            # Инициализируем кампании в порядке из campaigns.json
            for campaign_data in campaigns_list_from_json:
                campaign_id = campaign_data.get('Id')
                campaign_name = campaign_data.get('Name', '')
                # Получаем project_id из mapping (из БД), если нет - используем 0
                project_id = project_mapping.get(campaign_id, 0)
                
                if campaign_id:
                    # Определяем категорию по названию кампании (регистр не учитывается)
                    campaign_name_lower = campaign_name.lower()
                    category = "РСЯ" if "рся" in campaign_name_lower else "Поиск" if "поиск" in campaign_name_lower else "Неопределено"
                    
                    # Определяем тип кампании: "ключи" или "интересы"
                    campaign_type = "ключи" if campaign_id in campaigns_with_keywords else "интересы"
                    
                    campaigns_dict[campaign_id] = {
                        'campaign_id': campaign_id,
                        'campaign_name': campaign_name,
                        'category': category,
                        'campaign_type': campaign_type,
                        'project_id': project_id,
                        'ads': [],
                        'title_text_combinations': [],  # Изменили на список для сохранения порядка
                        'callouts': [],  # Изменили на список для сохранения порядка
                        'landing_pages': set(),  # Посадочные страницы (Href)
                        'sitelinks_titles': [],  # Изменили на список для сохранения порядка
                        'sitelinks_descriptions': []  # Изменили на список для сохранения порядка
                    }
            
            # Теперь обрабатываем объявления и заполняем данные кампаний
            for ad in ads:
                campaign_id = ad.get('CampaignId')
                if campaign_id in campaigns_dict:
                    campaigns_dict[campaign_id]['ads'].append(ad)
                
                # Собираем уникальные комбинации заголовок + текст
                if ad.get('Type') == 'TEXT_AD' and ad.get('TextAd'):
                    text_ad = ad['TextAd']
                    title = text_ad.get('Title')
                    text = text_ad.get('Text')
                    href = text_ad.get('Href')
                    sitelink_set_id = text_ad.get('SitelinkSetId')
                    
                    # Добавляем только если есть И заголовок И текст
                    if title and text:
                        # Создаем уникальную комбинацию
                        combination = f"{title} | {text}"
                        if combination not in campaigns_dict[campaign_id]['title_text_combinations']:
                            campaigns_dict[campaign_id]['title_text_combinations'].append(combination)
                    
                    # Посадочные страницы (Href) - для заголовка кампании
                    if href:
                        campaigns_dict[campaign_id]['landing_pages'].add(href)
                    
                    # Быстрые ссылки из SitelinkSetId
                    if sitelink_set_id and sitelink_set_id in sitelinks_dict:
                        for sitelink in sitelinks_dict[sitelink_set_id]:
                            if sitelink.get('Title') and sitelink['Title'] not in campaigns_dict[campaign_id]['sitelinks_titles']:
                                campaigns_dict[campaign_id]['sitelinks_titles'].append(sitelink['Title'])
                            if sitelink.get('Description') and sitelink['Description'] not in campaigns_dict[campaign_id]['sitelinks_descriptions']:
                                campaigns_dict[campaign_id]['sitelinks_descriptions'].append(sitelink['Description'])
                    
                    # Обрабатываем AdExtensions для получения уникальных уточнений
                    ad_extensions = text_ad.get('AdExtensions', [])
                    for ad_ext in ad_extensions:
                        if ad_ext.get('Type') == 'CALLOUT':
                            ext_id = ad_ext.get('AdExtensionId')
                            if ext_id and ext_id in extensions_dict:
                                callout_text = extensions_dict[ext_id]
                                if callout_text not in campaigns_dict[campaign_id]['callouts']:
                                    campaigns_dict[campaign_id]['callouts'].append(callout_text)
            
            # Группируем кампании по project_id
            campaigns_by_project = {}
            for campaign_data in campaigns_list_from_json:
                campaign_id = campaign_data.get('Id')
                if campaign_id and campaign_id in campaigns_dict:
                    data = campaigns_dict[campaign_id]
                    project_id = data['project_id']
                    
                    if project_id not in campaigns_by_project:
                        campaigns_by_project[project_id] = []
                    
                    campaigns_by_project[project_id].append({
                        'campaign_id': campaign_id,
                        'campaign_name': data['campaign_name'],
                        'category': data['category'],
                        'campaign_type': data['campaign_type'],
                        'project_id': data['project_id'],
                        'ads': data['ads'],
                        'title_text_combinations': data['title_text_combinations'],
                        'callouts': data['callouts'],
                        'landing_pages': list(data['landing_pages']),
                        'sitelinks_titles': data['sitelinks_titles'],
                        'sitelinks_descriptions': data['sitelinks_descriptions']
                    })
            
            # Формируем итоговый список: сначала по проектам, внутри каждого проекта - сначала РСЯ, потом Поиск
            campaigns_list = []
            # Сортируем project_id как числа (по возрастанию)
            for project_id in sorted(campaigns_by_project.keys(), key=lambda x: int(x) if isinstance(x, (int, str)) and str(x).isdigit() else 999999):
                project_campaigns = campaigns_by_project[project_id]
                
                # Сначала добавляем РСЯ кампании этого проекта
                for campaign in project_campaigns:
                    if campaign['category'] == "РСЯ":
                        campaigns_list.append(campaign)
                
                # Затем добавляем Поиск кампании этого проекта
                for campaign in project_campaigns:
                    if campaign['category'] == "Поиск":
                        campaigns_list.append(campaign)
            
            # Подсчитываем кампании по типам
            keyword_campaigns = [c for c in campaigns_list if c['campaign_type'] == 'ключи']
            interest_campaigns = [c for c in campaigns_list if c['campaign_type'] == 'интересы']
            
            print(f"✅ Найдено кампаний: {len(campaigns_list)} (ключи: {len(keyword_campaigns)}, интересы: {len(interest_campaigns)})")
            return campaigns_list
            
        except Exception as e:
            print(f"❌ Ошибка при получении данных о кампаниях: {e}")
            import traceback
            traceback.print_exc()
            return []

    def create_section_6(self, doc: Document, report_data: Dict) -> None:
        """Создать пятый раздел отчета"""
        # Получаем тексты из БД
        report_13 = self.get_report_text('report_13')
        report_14 = self.get_report_text('report_14')
        report_15 = self.get_report_text('report_15')
        report_16 = self.get_report_text('report_16')
        report_17 = self.get_report_text('report_17')
        report_18 = self.get_report_text('report_18')
        report_19 = self.get_report_text('report_19')
        report_20 = self.get_report_text('report_20')
        
        if not all([report_13, report_14, report_15, report_16, report_17, report_18, report_19, report_20]):
            print("❌ Не удалось получить все необходимые тексты для раздела 5")
            return

        # 1. report_13 + digital_project
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # report_13 жирным
        run1 = p.add_run(report_13)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # digital_project обычным
        run2 = p.add_run(f" {report_data['digital_project'] or ''}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. report_14 + start_date + end_date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        start_date = report_data['start_date']
        end_date = report_data['end_date']
        # report_14 жирным
        run1 = p.add_run(report_14)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # даты обычным
        if start_date and end_date:
            start_formatted = start_date.strftime("%d.%m.%Y")
            end_formatted = end_date.strftime("%d.%m.%Y")
            run2 = p.add_run(f" с {start_formatted} по {end_formatted}")
        else:
            run2 = p.add_run("")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_15 + goals (с новой строки)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # report_15 жирным
        run1 = p.add_run(report_15)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # goals обычным с новой строки
        run2 = p.add_run(f"\n{report_data['goals'] or ''}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 4. report_16 + tasks (с новой строки)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # report_16 жирным
        run1 = p.add_run(report_16)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # tasks обычным с новой строки
        run2 = p.add_run(f"\n{report_data['tasks'] or ''}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

        # 5. report_17 + description_target_audience (с новой строки)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # report_17 жирным
        run1 = p.add_run(report_17)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # description_target_audience обычным с новой строки
        run2 = p.add_run(f"\n{report_data['description_target_audience'] or ''}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

        # 6. report_18 + requirements_visual_materials
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # report_18 жирным
        run1 = p.add_run(report_18)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # requirements_visual_materials обычным
        run2 = p.add_run(f" {report_data['requirements_visual_materials'] or ''}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

        # 7. report_19 + requirements_text_materials
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # report_19 жирным
        run1 = p.add_run(report_19)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # requirements_text_materials обычным
        run2 = p.add_run(f" {report_data['requirements_text_materials'] or ''}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

        # 8. report_20 + KPI шаблон
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        clicks = self.format_number_with_spaces(report_data['kpi_plan_clicks'])
        reject = self.format_percentage(report_data['kpi_plan_reject'])
        # report_20 жирным
        run1 = p.add_run(report_20)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        # KPI шаблон обычным
        run2 = p.add_run(f" не менее {clicks} кликов, процент отказа не более {reject} %")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

    def create_section_7(self, doc: Document, report_data: Dict) -> None:
        """Создать шестой раздел отчета"""
        import tempfile
        import os
        
        # Получаем тексты из БД
        report_22 = self.get_report_text('report_22')
        report_23 = self.get_report_text('report_23')
        report_24 = self.get_report_text('report_24')
        report_25 = self.get_report_text('report_25')
        report_26 = self.get_report_text('report_26')
        
        if not all([report_22, report_23, report_24, report_25, report_26]):
            print("❌ Не удалось получить все необходимые тексты для раздела 6")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"{report_22} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_23 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_23} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 1 интервал (исключение)
        doc.add_paragraph()

        # Создаем таблицу переписки
        if report_data['correspondence']:
            # Создаем таблицу с заголовками
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'  # Стиль с видимыми границами
            table.allow_autofit = False
            
            # Настраиваем ширину таблицы
            table.width = Inches(6.0)  # Примерная ширина страницы A4
            # Колонки: 20%, 20%, 60%
            table.columns[0].width = Inches(1.2)  # 20% от 6 inches
            table.columns[1].width = Inches(1.2)  # 20% от 6 inches
            table.columns[2].width = Inches(3.6)  # 60% от 6 inches

            # Заголовки таблицы
            header_cells = table.rows[0].cells
            
            # Первая колонка (report_24)
            header_cells[0].text = report_24
            header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            header_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            header_cells[0].paragraphs[0].runs[0].bold = True
            
            # Вторая колонка (report_25)
            header_cells[1].text = report_25
            header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
            header_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
            header_cells[1].paragraphs[0].runs[0].bold = True
            
            # Третья колонка (report_26)
            header_cells[2].text = report_26
            header_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
            header_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
            header_cells[2].paragraphs[0].runs[0].bold = True

            # Добавляем данные переписки
            for correspondence_item in report_data['correspondence']:
                row = table.add_row()
                cells = row.cells
                
                # Тема письма
                cells[0].text = correspondence_item[4] or ''  # theme
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                cells[0].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # Дата отправки
                if correspondence_item[2]:  # date_sent
                    date_formatted = correspondence_item[2].strftime("%d.%m.%Y")
                else:
                    date_formatted = ""
                cells[1].text = date_formatted
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                cells[1].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # Изображение (если есть ссылка)
                if correspondence_item[3]:  # file_link
                    try:
                        # Загружаем изображение из MinIO (silent=True, чтобы не выводить ошибку для несуществующих файлов)
                        image_data = self.load_image_from_minio(correspondence_item[3], silent=True)
                        
                        if image_data:
                            # Очищаем ячейку от текста
                            cells[2].text = ""
                            
                            # Добавляем изображение в ячейку
                            paragraph = cells[2].paragraphs[0]
                            run = paragraph.add_run()
                            
                            # Вставляем изображение
                            
                            # Создаем временный файл для изображения
                            
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                                temp_file.write(image_data)
                                temp_file_path = temp_file.name
                            
                            try:
                                # Определяем размеры изображения
                                with Image.open(temp_file_path) as img:
                                    original_width, original_height = img.size
                                    
                                    # Максимальная ширина 350px
                                    max_width_px = 350
                                    
                                    # Рассчитываем пропорциональные размеры
                                    if original_width > max_width_px:
                                        # Масштабируем по ширине
                                        scale_factor = max_width_px / original_width
                                        new_width = max_width_px
                                        new_height = int(original_height * scale_factor)
                                        
                                        # Конвертируем в дюймы
                                        width_inches = Inches(new_width / 96)  # 96 DPI
                                        height_inches = Inches(new_height / 96)
                                    else:
                                        # Изображение уже подходящего размера
                                        width_inches = Inches(original_width / 96)
                                        height_inches = Inches(original_height / 96)
                                    
                                    # Добавляем изображение с рассчитанными размерами
                                    run.add_picture(temp_file_path, width=width_inches, height=height_inches)
                                
                                # Выравниваем изображение по центру
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                            finally:
                                # Удаляем временный файл
                                os.unlink(temp_file_path)
                        else:
                            cells[2].text = f"[Ошибка загрузки: {correspondence_item[3]}]"
                            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                            cells[2].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            
                    except Exception as e:
                        cells[2].text = f"[Ошибка обработки изображения: {str(e)}]"
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                        cells[2].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                else:
                    cells[2].text = ""
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                    cells[2].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE



    def create_section_8(self, doc: Document, report_data: Dict) -> None:
        """Создать седьмой раздел отчета"""
        # Получаем тексты из БД
        report_27 = self.get_report_text('report_27')
        report_28 = self.get_report_text('report_28')
        report_29 = self.get_report_text('report_29')
        report_30 = self.get_report_text('report_30')
        report_31 = self.get_report_text('report_31')
        report_32 = self.get_report_text('report_32')
        report_38 = self.get_report_text('report_38')
        
        if not all([report_27, report_28, report_29, report_30, report_31, report_32, report_38]):
            print("❌ Не удалось получить все необходимые тексты для раздела 7")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"{report_27} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_38 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_38} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Пустая строка перед таблицей
        doc.add_paragraph()

        # Создаем таблицу с данными (без заголовков)
        table = doc.add_table(rows=5, cols=2)  # 5 строк для данных
        table.style = 'Table Grid'  # Стиль с видимыми границами
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы
        table.width = Inches(6.0)  # Примерная ширина страницы A4
        # Колонки: 30%, 70%
        table.columns[0].width = Inches(1.8)  # 30% от 6 inches
        table.columns[1].width = Inches(4.2)  # 70% от 6 inches

        # Добавляем данные в таблицу
        # 1. digital_project
        cells = table.rows[0].cells
        cells[0].text = report_28
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True
        
        cells[1].text = report_data['digital_project'] or ''
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # 2. list_recommended_campaign_types
        cells = table.rows[1].cells
        cells[0].text = report_29
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True
        
        cells[1].text = report_data['list_recommended_campaign_types'] or ''
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # 3. list_recommended_formats_ads
        cells = table.rows[2].cells
        cells[0].text = report_30
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True
        
        cells[1].text = report_data['list_recommended_formats_ads'] or ''
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # 4. description_target_audience
        cells = table.rows[3].cells
        cells[0].text = report_31
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True
        
        cells[1].text = report_data['description_target_audience_requests'] or ''
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # 5. interests
        cells = table.rows[4].cells
        cells[0].text = report_32
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True
        
        cells[1].text = report_data['interests'] or ''
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # Добавляем таблицу с ключевыми фразами и минус-словами
        try:
            # Получаем тексты из БД
            report_33 = self.get_report_text('report_33')
            report_34 = self.get_report_text('report_34')
            report_35 = self.get_report_text('report_35')
            report_36 = self.get_report_text('report_36')
            report_37 = self.get_report_text('report_37')
            
            if not all([report_33, report_34, report_35, report_36, report_37]):
                print("❌ Не удалось получить все необходимые тексты для таблицы ключевых фраз")
                return

            # Получаем mapping campaign_id -> project_id из БД
            project_mapping = self.get_campaign_project_mapping(report_data.get('campany_yandex_direct'))

            # Получаем данные кампаний
            campaigns_data = self.get_campaigns_data_for_report(report_data['id'], project_mapping)
            if not campaigns_data:
                print("❌ Не удалось получить данные кампаний")
                return

            # Загружаем ключевые фразы
            keywords_data = self.load_file_from_minio(report_data['id'], f"keywords_traffic_forecast_{report_data['id']}.json")
            if not keywords_data:
                print("❌ Не удалось загрузить данные о ключевых фразах")
                return

            # Загружаем данные кампаний для минус-слов
            campaigns_json_data = self.load_file_from_minio(report_data['id'], "campaigns.json")
            if not campaigns_json_data:
                print("❌ Не удалось загрузить данные кампаний для минус-слов")
                return

            # Определяем кампании "ключи"
            campaigns_with_keywords = set()
            keywords = keywords_data.get('result', {}).get('Keywords', [])
            for keyword in keywords:
                campaign_id = keyword.get('CampaignId')
                keyword_text = keyword.get('Keyword', '')
                
                # Исключаем ключевые фразы с текстом "---autotargeting"
                if campaign_id and keyword_text != "---autotargeting":
                    campaigns_with_keywords.add(campaign_id)

            # Фильтруем только кампании "ключи", сохраняя исходный порядок
            keyword_campaigns_list = [c for c in campaigns_data if c.get('campaign_type') == 'ключи']
            
            if not keyword_campaigns_list:
                print("❌ Не найдено кампаний типа 'ключи'")
                return

            # Пустая строка перед таблицей
            doc.add_paragraph()

            # Создаем таблицу с одной колонкой
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.allow_autofit = False
            table.width = Inches(6.0)

            # Первая строка - report_33 жирная по центру
            first_cell = table.rows[0].cells[0]
            first_cell.text = report_33
            first_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            first_cell.paragraphs[0].runs[0].font.size = Pt(12)
            first_cell.paragraphs[0].runs[0].bold = True
            first_cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            first_cell.paragraphs[0].paragraph_format.line_spacing = Pt(13.8)  # 12pt * 1.15

            # Список для сбора всех минус-слов из всех кампаний
            all_negative_keywords = []

            # Счетчик для сквозной нумерации кампаний
            campaign_counter = 1

            # Обрабатываем кампании в том порядке, в котором они пришли из get_campaigns_data_for_report
            for campaign in keyword_campaigns_list:
                campaign_id = campaign['campaign_id']
                
                # Вторая строка - шаблон с номером кампании и ссылками
                row = table.add_row()
                cell = row.cells[0]
                
                # Получаем ссылки на проект из кампании
                landing_pages = campaign.get('landing_pages', [])
                links_text = ', '.join(landing_pages) if landing_pages else ''
                
                cell_text = f"{report_34} {campaign_counter} {report_35} ({links_text}) {report_36}"
                campaign_counter += 1
                cell.text = cell_text
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                cell.paragraphs[0].paragraph_format.line_spacing = Pt(13.8)  # 12pt * 1.15

                # Третья строка - ключевые фразы (каждая в отдельной строке)
                # Фильтруем ключевые фразы для данной кампании
                campaign_keywords = []
                for keyword in keywords:
                    if (keyword.get('CampaignId') == campaign_id and 
                        keyword.get('Keyword', '') != "---autotargeting"):
                        keyword_text = keyword.get('Keyword', '')
                        if keyword_text:
                            campaign_keywords.append(keyword_text)

                # Добавляем каждую ключевую фразу в отдельную строку
                for keyword_text in campaign_keywords:
                    row = table.add_row()
                    cell = row.cells[0]
                    # Убираем кавычки из ключевых фраз
                    clean_keyword = keyword_text.strip('"')
                    cell.text = clean_keyword
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    cell.paragraphs[0].paragraph_format.line_spacing = Pt(13.8)  # 12pt * 1.15

                # Ищем минус-слова для данной кампании и добавляем в общий список
                negative_keywords = []
                campaigns_list = campaigns_json_data.get('result', {}).get('Campaigns', [])
                for campaign_data in campaigns_list:
                    if campaign_data.get('Id') == campaign_id:
                        negative_data = campaign_data.get('NegativeKeywords')
                        if negative_data and negative_data.get('Items'):
                            negative_keywords = negative_data['Items']
                        break

                # Собираем все минус-слова в общий список
                if negative_keywords:
                    all_negative_keywords.extend(negative_keywords)

            # После обработки всех кампаний добавляем заголовок минус-слов всегда
            # Добавляем заголовок report_37
            row = table.add_row()
            cell = row.cells[0]
            cell.text = report_37
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(13.8)  # 12pt * 1.15

            # Добавляем минус-слова или прочерк, если их нет
            if all_negative_keywords:
                # Добавляем все минус-слова построчно (каждое в отдельной строке)
                for negative_keyword in all_negative_keywords:
                    row = table.add_row()
                    cell = row.cells[0]
                    # Убираем кавычки из минус-слов
                    clean_negative = negative_keyword.strip('"')
                    cell.text = clean_negative
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    cell.paragraphs[0].paragraph_format.line_spacing = Pt(13.8)  # 12pt * 1.15
            else:
                # Если минус-слов нет, добавляем прочерк
                row = table.add_row()
                cell = row.cells[0]
                cell.text = "-"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                cell.paragraphs[0].paragraph_format.line_spacing = Pt(13.8)  # 12pt * 1.15

        except Exception as e:
            print(f"❌ Ошибка при создании таблицы ключевых фраз: {e}")
            raise e

    def create_section_9(self, doc: Document, report_data: Dict, report_id: int) -> None:
        """Создать восьмой раздел отчета с изображениями"""
        import tempfile
        import requests
        
        # Получаем тексты из БД
        report_39 = self.get_report_text('report_39')
        report_40 = self.get_report_text('report_40')
        report_41 = self.get_report_text('report_41')
        
        if not all([report_39, report_40, report_41]):
            print("❌ Не удалось получить все необходимые тексты для раздела 8")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"{report_39} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_40 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_40} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Пустая строка после номера контракта
        doc.add_paragraph()

        # 4. report_41 (жирным и подчеркнутым)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_41)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.underline = True

        # Получаем уникальные изображения
        unique_images = self.get_unique_images_for_report(report_id)
        
        if not unique_images:
            print("⚠ Нет изображений для отображения")
            return

        # Создаем таблицу с изображениями
        table = doc.add_table(rows=len(unique_images), cols=2)
        table.style = 'Table Grid'  # Стиль с видимыми границами
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы (уменьшена на 30% от 6.0)
        table.width = Inches(4.2)
        
        # Устанавливаем фиксированную ширину колонок: 20%, 80%
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Для первой колонки - 20% (0.84 inches = 1210 twips)
        for cell in table.columns[0].cells:
            cell.width = Inches(0.84)
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '1210')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        # Для второй колонки - 80% (3.36 inches = 4838 twips)
        for cell in table.columns[1].cells:
            cell.width = Inches(3.36)
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '4838')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # Заполняем таблицу
        for i, image in enumerate(unique_images, 1):
            cells = table.rows[i-1].cells
            
            # Первая колонка - порядковый номер
            cells[0].text = str(i)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            cells[0].paragraphs[0].runs[0].bold = True
            # Центрируем по вертикали
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Вторая колонка - изображение
            try:
                # Загружаем изображение по URL
                # Хардкод response = requests.get(image['original_url'], timeout=30) Preview
                response = requests.get(image['preview_url'], timeout=30)
                if response.status_code == 200:
                    # Создаем временный файл
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        # Определяем размеры изображения
                        with Image.open(temp_file_path) as img:
                            original_width, original_height = img.size
                            
                            # Максимальная высота 100px
                            max_height_px = 100
                            # Минимальная высота для квадратных изображений 80px
                            min_height_px = 80
                            
                            # Рассчитываем пропорциональные размеры
                            if original_height > max_height_px:
                                # Масштабируем по высоте
                                scale_factor = max_height_px / original_height
                                new_width = int(original_width * scale_factor)
                                new_height = max_height_px
                            elif original_height < min_height_px and original_width == original_height:
                                # Для квадратных изображений увеличиваем до минимальной высоты
                                scale_factor = min_height_px / original_height
                                new_width = int(original_width * scale_factor)
                                new_height = min_height_px
                            else:
                                # Изображение уже подходящего размера
                                new_width = original_width
                                new_height = original_height
                            
                            # Конвертируем в дюймы
                            width_inches = Inches(new_width / 96)  # 96 DPI
                            height_inches = Inches(new_height / 96)
                            
                            # Добавляем изображение в ячейку
                            paragraph = cells[1].paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(temp_file_path, width=width_inches, height=height_inches)
                            
                            # Выравниваем изображение по центру
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                    finally:
                        # Удаляем временный файл
                        os.unlink(temp_file_path)
                else:
                    cells[1].text = f"[Ошибка загрузки изображения: {response.status_code}]"
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                    
            except Exception as e:
                cells[1].text = f"[Ошибка обработки изображения: {str(e)}]"
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # Получаем текст report_51
        report_51 = self.get_report_text('report_51')
        if report_51:
            # Одна пустая строка после таблицы
            doc.add_paragraph()
            
            # report_51 слева жирным
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(report_51)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            
            # Еще одна пустая строка
            doc.add_paragraph()

        # Создаем таблицу с данными о кампаниях
        self.create_campaigns_table(doc, report_data, report_id)

    def create_campaigns_table(self, doc: Document, report_data: Dict, report_id: int) -> None:
        """Создать таблицу с данными о кампаниях"""
        # Получаем тексты из БД
        report_43 = self.get_report_text('report_43')
        report_44 = self.get_report_text('report_44')
        report_45 = self.get_report_text('report_45')
        report_46 = self.get_report_text('report_46')
        report_47 = self.get_report_text('report_47')
        report_48 = self.get_report_text('report_48')
        report_49 = self.get_report_text('report_49')
        report_50 = self.get_report_text('report_50')
        report_52 = self.get_report_text('report_52')
        
        if not all([report_43, report_44, report_45, report_46, report_47, report_48, report_49, report_50, report_52]):
            print("❌ Не удалось получить все необходимые тексты для таблицы кампаний")
            return

        # Получаем mapping campaign_id -> project_id из БД
        project_mapping = self.get_campaign_project_mapping(report_data.get('campany_yandex_direct'))

        # Получаем данные о кампаниях
        campaigns_data = self.get_campaigns_data_for_report(report_id, project_mapping)
        if not campaigns_data:
            print("⚠ Нет данных о кампаниях для отображения")
            return

        # Создаем таблицу
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'  # Стиль с видимыми границами
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы
        table.width = Inches(6.0)  # Примерная ширина страницы A4
        # Колонки: равномерно распределены
        for i in range(5):
            table.columns[i].width = Inches(1.2)  # 20% каждая

        # Заголовки таблицы
        header_cells = table.rows[0].cells
        headers = [report_43, report_44, report_45, report_46, report_47]
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Фильтруем только кампании "ключи", сохраняя исходный порядок
        keyword_campaigns_list = [c for c in campaigns_data if c.get('campaign_type') == 'ключи']
        interest_campaigns_list = [c for c in campaigns_data if c.get('campaign_type') == 'интересы']
        
        print(f"📊 Кампании 'ключи': {len(keyword_campaigns_list)}, 'интересы': {len(interest_campaigns_list)} (пропускаются)")
        
        # Счетчик для сквозной нумерации кампаний
        campaign_counter = 1
        
        # Обрабатываем кампании в том порядке, в котором они пришли из get_campaigns_data_for_report
        for campaign in keyword_campaigns_list:
            # Создаем объединенную ячейку для заголовка кампании
            header_row = table.add_row()
            
            # Заголовок кампании - объединяем все 5 ячеек
            campaign_header = f"{report_48} {campaign_counter} {report_49} ({', '.join(campaign['landing_pages'])}) {report_50}"
            campaign_counter += 1
            header_row.cells[0].text = campaign_header
            header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            header_row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            header_row.cells[0].paragraphs[0].runs[0].bold = True
            
            # Объединяем ячейки (остальные 4 ячейки объединяем с первой)
            # Используем метод merge() для объединения ячеек
            merged_cell = header_row.cells[0].merge(header_row.cells[4])

            # Добавляем строки данных для кампании
            self._add_campaign_data_rows(table, campaign)
        
        # Обрабатываем кампании типа "интересы" - ЗАКОММЕНТИРОВАНО
        # for campaign_num, campaign in enumerate(interest_campaigns, 1):
        #     # Создаем объединенную ячейку для заголовка кампании
        #     header_row = table.add_row()
        #     
        #     # Заголовок кампании - объединяем все 5 ячеек (используем report_52)
        #     campaign_header = f"{report_52} {campaign_num} {report_49} ({', '.join(campaign['landing_pages'])}) {report_50}"
        #     header_row.cells[0].text = campaign_header
        #     header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     header_row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        #     header_row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        #     header_row.cells[0].paragraphs[0].runs[0].bold = True
        #     
        #     # Объединяем ячейки (остальные 4 ячейки объединяем с первой)
        #     # Используем метод merge() для объединения ячеек
        #     merged_cell = header_row.cells[0].merge(header_row.cells[4])
        #     
        #     # Добавляем строки данных для кампании
        #     self._add_campaign_data_rows(table, campaign)

    def _add_campaign_data_rows(self, table, campaign):
        """Добавить строки данных для кампании"""
        # Получаем тексты из БД для форматирования
        # report_43 = self.get_report_text('report_43')
        # report_44 = self.get_report_text('report_44')
        # report_45 = self.get_report_text('report_45')
        # report_46 = self.get_report_text('report_46')
        # report_47 = self.get_report_text('report_47')

        # Находим максимальное количество строк для данной кампании
        # Учитываем только столбцы 1, 2, 4, 5 (исключаем столбец 3 - уточнения)
        max_rows = max(
            len(campaign['title_text_combinations']),
            len(campaign['sitelinks_titles']),
            len(campaign['sitelinks_descriptions'])
        )

        # Создаем строки данных только если есть хотя бы одно значение
        if max_rows > 0:
            for i in range(max_rows):
                data_row = table.add_row()
                cells = data_row.cells
                
                # Комбинации заголовок + текст (report_43 и report_44)
                if i < len(campaign['title_text_combinations']):
                    combination = campaign['title_text_combinations'][i]
                    # Разделяем комбинацию на заголовок и текст
                    if ' | ' in combination:
                        title, text = combination.split(' | ', 1)
                        cells[0].text = title  # report_43
                        cells[1].text = text   # report_44
                    else:
                        cells[0].text = combination
                        cells[1].text = "—"
                else:
                    cells[0].text = "—"
                    cells[1].text = "—"
                
                # Форматирование для заголовков (report_43)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                cells[0].paragraphs[0].paragraph_format.line_spacing = 1.0
                
                # Форматирование для текстов (report_44)
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                cells[1].paragraphs[0].paragraph_format.line_spacing = 1.0
                
                # Уточнения (report_45) - только в первой строке
                if i == 0 and campaign['callouts']:
                    # Объединяем все уточнения в одну ячейку с пустыми строками между ними
                    all_callouts = '\n\n'.join(campaign['callouts'])
                    cells[2].text = all_callouts
                elif i == 0:
                    cells[2].text = "—"
                else:
                    # Очищаем ячейку для объединения
                    cells[2].text = ""
                
                # Форматирование для уточнений (report_45)
                if i == 0:
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                    cells[2].paragraphs[0].paragraph_format.line_spacing = 1.0
                
                # Заголовки быстрых ссылок (report_46)
                if i < len(campaign['sitelinks_titles']):
                    cells[3].text = campaign['sitelinks_titles'][i]
                else:
                    cells[3].text = "—"
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[3].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[3].paragraphs[0].runs[0].font.size = Pt(12)
                cells[3].paragraphs[0].paragraph_format.line_spacing = 1.0
                
                # Описания быстрых ссылок (report_47)
                if i < len(campaign['sitelinks_descriptions']):
                    cells[4].text = campaign['sitelinks_descriptions'][i]
                else:
                    cells[4].text = "—"
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[4].paragraphs[0].runs[0].font.size = Pt(12)
                cells[4].paragraphs[0].paragraph_format.line_spacing = 1.0
        
            # Объединяем ячейки уточнений по вертикали
            if max_rows > 1 and campaign['callouts']:
                # Находим строки данных для данной кампании
                data_rows_start = len(table.rows) - max_rows
                # Объединяем ячейки уточнений (колонка 2) по вертикали
                for i in range(1, max_rows):
                    row_index = data_rows_start + i
                    if row_index < len(table.rows):
                        # Объединяем ячейку с предыдущей строкой
                        table.rows[data_rows_start].cells[2].merge(table.rows[row_index].cells[2])

    def create_section_9_new(self, doc: Document, report_data: Dict) -> None:
        """Создать девятый раздел отчета"""
        # Получаем тексты из БД
        report_53 = self.get_report_text('report_53')
        report_54 = self.get_report_text('report_54')
        
        if not all([report_53, report_54]):
            print("❌ Не удалось получить все необходимые тексты для раздела 9")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня: report_53 + date_request + № + application_number
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"{report_53} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_54 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_54} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 4. Одна пустая строка
        doc.add_paragraph()

        # 5. examples_published_ads слева
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_data['examples_published_ads'] or '')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def load_campaign_stats_summary(self, report_id: int) -> Optional[Dict]:
        """Загрузить данные из campaign_stats_summary файла из MinIO"""
        try:
            # Формируем имя файла
            filename = f"campaign_stats_summary_{report_id}.json"
            return self.load_file_from_minio(report_id, filename)
        except Exception as e:
            print(f"❌ Ошибка при загрузке campaign_stats_summary: {e}")
            # return None
            raise e

    def load_campaign_stats(self, report_id: int) -> Optional[Dict]:
        """Загрузить данные из campaign_stats файла из MinIO"""
        try:
            # Формируем имя файла
            filename = f"campaign_stats_{report_id}.json"
            return self.load_file_from_minio(report_id, filename)
        except Exception as e:
            print(f"❌ Ошибка при загрузке campaign_stats: {e}")
            # return None
            raise e

    def load_adgroup_stats(self, report_id: int) -> Optional[Dict]:
        """Загрузить данные из adgroup_stats файла из MinIO"""
        try:
            # Формируем имя файла
            filename = f"adgroup_stats_{report_id}.json"
            return self.load_file_from_minio(report_id, filename)
        except Exception as e:
            print(f"❌ Ошибка при загрузке adgroup_stats: {e}")
            # return None
            raise e

    def load_ad_stats(self, report_id: int) -> Optional[Dict]:
        """Загрузить данные из ad_stats файла из MinIO"""
        try:
            # Формируем имя файла
            filename = f"ad_stats_{report_id}.json"
            return self.load_file_from_minio(report_id, filename)
        except Exception as e:
            print(f"❌ Ошибка при загрузке ad_stats: {e}")
            # return None
            raise e

    def load_ads_report(self, report_id: int) -> Optional[Dict]:
        """Загрузить данные из ads_report файла из MinIO"""
        try:
            # Формируем имя файла
            filename = f"ads_report_{report_id}.json"
            return self.load_file_from_minio(report_id, filename)
        except Exception as e:
            print(f"❌ Ошибка при загрузке ads_report: {e}")
            # return None
            raise e

    def load_campaigns(self, report_id: int) -> Optional[Dict]:
        """Загрузить данные из campaigns файла из MinIO"""
        try:
            # Формируем имя файла (без префикса номера отчета)
            filename = "campaigns.json"
            return self.load_file_from_minio(report_id, filename)
        except Exception as e:
            print(f"❌ Ошибка при загрузке campaigns: {e}")
            # return None
            raise e

    def get_top_ads(self, report_id: int) -> List[Dict]:
        """Получить топ объявлений по кликам с фильтрацией по BounceRate"""
        try:
            # Загружаем данные
            ad_stats_data = self.load_ad_stats(report_id)
            ads_report_data = self.load_ads_report(report_id)
            campaigns_data = self.load_campaigns(report_id)
            
            if not ad_stats_data or not ads_report_data or not campaigns_data:
                print("❌ Не удалось загрузить данные для топ объявлений")
                return []
            
            # Получаем объявления из ad_stats
            ad_stats_rows = ad_stats_data.get('result', {}).get('rows', [])
            if not ad_stats_rows:
                print("❌ Нет данных объявлений в ad_stats")
                return []
            
            # Создаем словари для быстрого поиска
            ads_report_dict = {}
            for ad in ads_report_data.get('result', {}).get('Ads', []):
                ads_report_dict[ad['Id']] = ad
            
            campaigns_dict = {}
            for campaign in campaigns_data.get('result', {}).get('Campaigns', []):
                campaigns_dict[campaign['Id']] = campaign
            
            # Фильтруем объявления по BounceRate < 35%
            filtered_ads = []
            for ad_stat in ad_stats_rows:
                bounce_rate = ad_stat.get('BounceRate', 0)
                if bounce_rate < 35:
                    filtered_ads.append(ad_stat)
            
            # Если после фильтрации не осталось объявлений, убираем фильтр
            if not filtered_ads:
                print("⚠️ После фильтрации по BounceRate < 35% не осталось объявлений, убираем фильтр")
                filtered_ads = ad_stats_rows
            
            # Сортируем по количеству кликов (по убыванию)
            filtered_ads.sort(key=lambda x: x.get('Clicks', 0), reverse=True)
            
            # Берем топ N объявлений
            top_ads = filtered_ads[:TOP_ADS_COUNT]
            
            # Обогащаем данными из других файлов
            enriched_ads = []
            for ad_stat in top_ads:
                ad_id = ad_stat['AdId']
                campaign_id = ad_stat['CampaignId']
                
                # Получаем данные из ads_report
                ad_report = ads_report_dict.get(ad_id, {})
                
                # Получаем данные кампании
                campaign = campaigns_dict.get(campaign_id, {})
                campaign_name = campaign.get('Name', 'Неизвестная кампания')
                
                # Формируем ссылку на посадочную страницу
                href = ""
                if ad_report.get('TextAd'):
                    href = ad_report['TextAd'].get('Href', '')
                elif ad_report.get('ImageAd'):
                    href = ad_report['ImageAd'].get('Href', '')
                
                enriched_ad = {
                    'AdId': ad_id,
                    'CampaignId': campaign_id,
                    'CampaignName': campaign_name,
                    'Clicks': ad_stat.get('Clicks', 0),
                    'BounceRate': ad_stat.get('BounceRate', 0),
                    'AvgCpc': ad_stat.get('AvgCpc', 0),
                    'Href': href
                }
                enriched_ads.append(enriched_ad)
            
            return enriched_ads
            
        except Exception as e:
            print(f"❌ Ошибка при получении топ объявлений: {e}")
            # return []
            raise e

    def format_cost_from_kopecks(self, kopecks):
        """Конвертировать копейки в рубли с форматированием"""
        if kopecks is None:
            return "0"
        rubles = kopecks / 100.0
        return self.format_number_with_spaces(rubles)

    def create_section_10(self, doc: Document, report_data: Dict, report_id: int) -> None:
        """Создать десятый раздел отчета"""
        # Получаем тексты из БД
        report_55 = self.get_report_text('report_55')
        report_56 = self.get_report_text('report_56')
        report_54 = self.get_report_text('report_54')
        report_58 = self.get_report_text('report_58')
        report_59 = self.get_report_text('report_59')
        report_60 = self.get_report_text('report_60')
        report_61 = self.get_report_text('report_61')
        report_62 = self.get_report_text('report_62')
        report_63 = self.get_report_text('report_63')
        report_64 = self.get_report_text('report_64')
        report_65 = self.get_report_text('report_65')
        report_66 = self.get_report_text('report_66')
        report_67 = self.get_report_text('report_67')
        report_68 = self.get_report_text('report_68')
        report_69 = self.get_report_text('report_69')
        report_70 = self.get_report_text('report_70')
        report_71 = self.get_report_text('report_71')
        report_72 = self.get_report_text('report_72')
        report_73 = self.get_report_text('report_73')
        report_74 = self.get_report_text('report_74')
        report_75 = self.get_report_text('report_75')
        report_76 = self.get_report_text('report_76')
        report_77 = self.get_report_text('report_77')
        report_78 = self.get_report_text('report_78')
        report_79 = self.get_report_text('report_79')
        report_80 = self.get_report_text('report_80')
        report_81 = self.get_report_text('report_81')
        report_82 = self.get_report_text('report_82')
        report_83 = self.get_report_text('report_83')
        report_84 = self.get_report_text('report_84')
        report_85 = self.get_report_text('report_85')
        report_86 = self.get_report_text('report_86')
        report_87 = self.get_report_text('report_87')
        report_88 = self.get_report_text('report_88')
        report_89 = self.get_report_text('report_89')
        report_90 = self.get_report_text('report_90')
        report_91 = self.get_report_text('report_91')
        
        if not all([report_55, report_56, report_54, report_58, report_59, report_60, report_61, report_62, report_63, report_64, report_65, report_66, report_67, report_68, report_69, report_70, report_71, report_72, report_73, report_74, report_75, report_76, report_77, report_78, report_79, report_80, report_81, report_82, report_83, report_84, report_85, report_86, report_87, report_88, report_89, report_90, report_91]):
            print("❌ Не удалось получить все необходимые тексты для раздела 10")
            return

        # Загружаем данные из campaign_stats_summary
        campaign_stats = self.load_campaign_stats_summary(report_id)
        if not campaign_stats:
            print("❌ Не удалось загрузить данные campaign_stats_summary")
            return

        # Загружаем данные из campaign_stats
        campaign_stats_data = self.load_campaign_stats(report_id)
        if not campaign_stats_data:
            print("❌ Не удалось загрузить данные campaign_stats")
            return

        # Загружаем данные из adgroup_stats
        adgroup_stats_data = self.load_adgroup_stats(report_id)
        if not adgroup_stats_data:
            print("❌ Не удалось загрузить данные adgroup_stats")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня: report_55 + start_date + end_date + report_56 + date_request + № + application_number
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Форматируем даты периода
        start_date = report_data['start_date']
        end_date = report_data['end_date']
        if start_date and end_date:
            start_formatted = start_date.strftime("%d.%m.%Y")
            end_formatted = end_date.strftime("%d.%m.%Y")
            period_text = f"с {start_formatted} по {end_formatted} г."
        else:
            period_text = ""
        
        text = f"{report_55} {period_text} {report_56} {self.format_date(report_data['date_request'])} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_54 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_54} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 4. Пустая строка
        doc.add_paragraph()

        # 5. report_58 слева
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_58)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # 6. Создаем таблицу 3x3
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'  # Стиль с видимыми границами
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы
        table.width = Inches(6.0)  # Примерная ширина страницы A4
        # Колонки: равномерно распределены
        for i in range(3):
            table.columns[i].width = Inches(2.0)  # 33.33% каждая

        # Заполняем таблицу
        # 1 столбец
        # 1 строка - пустая
        cells = table.rows[0].cells
        cells[0].text = ""
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        
        # 2 строка - report_59 (жирным)
        cells = table.rows[1].cells
        cells[0].text = report_59
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True
        
        # 3 строка - report_60 (жирным)
        cells = table.rows[2].cells
        cells[0].text = report_60
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        cells[0].paragraphs[0].runs[0].bold = True

        # 2 столбец
        # 1 строка - report_61 (жирным)
        cells = table.rows[0].cells
        cells[1].text = report_61
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        cells[1].paragraphs[0].runs[0].bold = True
        
        # 2 строка - report_64 + kpi_plan_clicks (как числа, разряды отделяем пробелами, а десятных тут не пишем)
        cells = table.rows[1].cells
        clicks_formatted = self.format_number_with_spaces(report_data['kpi_plan_clicks'])
        cells[1].text = f"{report_64} {clicks_formatted}"
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        
        # 3 строка - Clicks из campaign_stats_summary
        cells = table.rows[2].cells
        summary_clicks = campaign_stats.get('summary', {}).get('Clicks', 0)
        clicks_summary_formatted = self.format_number_with_spaces(summary_clicks)
        cells[1].text = clicks_summary_formatted
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[1].paragraphs[0].runs[0].font.size = Pt(12)

        # 3 столбец
        # 1 строка - report_62 (жирным)
        cells = table.rows[0].cells
        cells[2].text = report_62
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[2].paragraphs[0].runs[0].font.size = Pt(12)
        cells[2].paragraphs[0].runs[0].bold = True
        
        # 2 строка - report_65 + kpi_plan_reject + "%"
        cells = table.rows[1].cells
        reject_formatted = self.format_percentage(report_data['kpi_plan_reject'])
        cells[2].text = f"{report_65} {reject_formatted}%"
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[2].paragraphs[0].runs[0].font.size = Pt(12)
        
        # 3 строка - BounceRate из campaign_stats_summary
        cells = table.rows[2].cells
        summary_bounce_rate = campaign_stats.get('summary', {}).get('BounceRate', 0)
        bounce_rate_formatted = self.format_number_with_spaces(summary_bounce_rate)
        cells[2].text = bounce_rate_formatted
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
        cells[2].paragraphs[0].runs[0].font.size = Pt(12)

        # 7. report_63 сразу после таблицы
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_63)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # 8. report_66 с новой строки
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_66)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # 9. Создаем таблицу с данными кампаний
        campaigns_rows = campaign_stats_data.get('result', {}).get('rows', [])
        if campaigns_rows:
            # Создаем таблицу с заголовками + данные + итоговая строка
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'  # Стиль с видимыми границами
            table.allow_autofit = False
            
            # Настраиваем ширину таблицы
            table.width = Inches(6.0)  # Примерная ширина страницы A4
            # Колонки согласно процентам: 5%, 30%, 15%, 15%, 15%, 10%, 10%
            table.columns[0].width = Inches(0.3)  # 5%
            table.columns[1].width = Inches(1.8)  # 30%
            table.columns[2].width = Inches(0.9)  # 15%
            table.columns[3].width = Inches(0.9)  # 15%
            table.columns[4].width = Inches(0.9)  # 15%
            table.columns[5].width = Inches(0.6)  # 10%
            table.columns[6].width = Inches(0.6)  # 10%

            # Заголовки таблицы
            header_cells = table.rows[0].cells
            headers = [report_67, report_68, report_69, report_70, report_71, report_72, report_73]
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Добавляем данные кампаний
            for i, campaign in enumerate(campaigns_rows, 1):
                row = table.add_row()
                cells = row.cells
                
                # 1. Номер (сквозная нумерация)
                cells[0].text = str(i)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                cells[0].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 2. CampaignName
                cells[1].text = campaign.get('CampaignName', '')
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                cells[1].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 3. CampaignId
                cells[2].text = str(campaign.get('CampaignId', ''))
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                cells[2].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 4. Impressions
                impressions = campaign.get('Impressions', 0)
                cells[3].text = self.format_number_with_spaces(impressions)
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[3].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[3].paragraphs[0].runs[0].font.size = Pt(12)
                cells[3].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 5. Clicks
                clicks = campaign.get('Clicks', 0)
                cells[4].text = self.format_number_with_spaces(clicks)
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[4].paragraphs[0].runs[0].font.size = Pt(12)
                cells[4].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 6. Ctr
                ctr = campaign.get('Ctr', 0)
                cells[5].text = self.format_number_with_spaces(ctr)
                cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[5].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[5].paragraphs[0].runs[0].font.size = Pt(12)
                cells[5].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 7. BounceRate
                bounce_rate = campaign.get('BounceRate', 0)
                cells[6].text = self.format_number_with_spaces(bounce_rate)
                cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[6].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[6].paragraphs[0].runs[0].font.size = Pt(12)
                cells[6].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Добавляем итоговую строку
            total_row = table.add_row()
            cells = total_row.cells
            
            # Объединяем первые 3 ячейки для report_74
            cells[0].text = report_74
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            cells[0].paragraphs[0].runs[0].bold = True
            cells[0].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # Объединяем ячейки 0, 1, 2
            merged_cell = cells[0].merge(cells[2])
            
            # 4. Impressions из summary
            summary_impressions = campaign_stats.get('summary', {}).get('Impressions', 0)
            cells[3].text = self.format_number_with_spaces(summary_impressions)
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[3].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[3].paragraphs[0].runs[0].font.size = Pt(12)
            cells[3].paragraphs[0].runs[0].bold = True
            cells[3].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # 5. Clicks из summary
            summary_clicks = campaign_stats.get('summary', {}).get('Clicks', 0)
            cells[4].text = self.format_number_with_spaces(summary_clicks)
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[4].paragraphs[0].runs[0].font.size = Pt(12)
            cells[4].paragraphs[0].runs[0].bold = True
            cells[4].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # 6. Ctr из summary
            summary_ctr = campaign_stats.get('summary', {}).get('Ctr', 0)
            cells[5].text = self.format_number_with_spaces(summary_ctr)
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[5].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[5].paragraphs[0].runs[0].font.size = Pt(12)
            cells[5].paragraphs[0].runs[0].bold = True
            cells[5].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # 7. BounceRate из summary
            summary_bounce_rate = campaign_stats.get('summary', {}).get('BounceRate', 0)
            cells[6].text = self.format_number_with_spaces(summary_bounce_rate)
            cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[6].paragraphs[0].runs[0].font.name = 'Times New Roman'
            cells[6].paragraphs[0].runs[0].font.size = Pt(12)
            cells[6].paragraphs[0].runs[0].bold = True
            cells[6].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 10. report_75 под таблицей слева жирным
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_75)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        # 11. Добавляем все изображения из папки url_1
        try:
            # Загружаем все скриншоты из папки url_1
            screenshot_counter = 1
            while True:
                # Формируем путь к изображению с нумерацией
                image_path = f"gen_report_context_contracts/data_yandex_direct/{report_id}_результаты/screenshots/url_1/screenshot_{screenshot_counter:03d}.png"
                
                # Загружаем изображение из MinIO (silent=True, чтобы не выводить ошибку для несуществующих файлов)
                image_data = self.load_image_from_minio(image_path, silent=True)
                
                if image_data:
                    # Создаем временный файл для изображения
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(image_data)
                        temp_file_path = temp_file.name
                    
                    try:
                        # Добавляем изображение в документ
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        # Вставляем изображение с максимальной шириной, Word автоматически подберет высоту
                        run.add_picture(temp_file_path, width=Inches(6.0))
                        
                        print(f"✅ Загружен скриншот: screenshot_{screenshot_counter:03d}.png")
                        
                    finally:
                        # Удаляем временный файл
                        os.unlink(temp_file_path)
                    
                    screenshot_counter += 1
                else:
                    # Если изображение не найдено, прерываем цикл
                    if screenshot_counter == 1:
                        print(f"❌ Не удалось загрузить первый скриншот: {image_path}")
                    break
                
        except Exception as e:
            print(f"❌ Ошибка при добавлении изображений: {e}")
            raise e

        # 12. Пустая строка
        doc.add_paragraph()

        # 13. report_76 слева жирным
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_76)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        # 14. Пустая строка
        doc.add_paragraph()

        # 14. Создаем таблицу с данными групп объявлений
        adgroup_rows = adgroup_stats_data.get('result', {}).get('rows', [])
        if adgroup_rows:
            # Создаем таблицу с заголовками + данные + итоговая строка
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'  # Стиль с видимыми границами
            table.allow_autofit = False
            
            # Настраиваем ширину таблицы (увеличиваем для реального расширения)
            total_width = Inches(13.0)
            table.width = total_width
            
            # Устанавливаем отрицательный отступ слева, чтобы таблица могла выходить за пределы текстового поля
            # Это позволит таблице занять больше места по ширине (и влево, и вправо)
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Устанавливаем отрицательный отступ влево (сдвиг таблицы)
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), str(-1080))  # -0.75 дюйма слева (1080 twips = 0.75 дюйма)
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)
            
            # Устанавливаем выравнивание таблицы по левому краю без ограничений
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'left')
            tblPr.append(jc)
            
            # Устанавливаем ширину таблицы как фиксированную (не автоподбор)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int(13.0 * 1440)))  # 13 дюймов в twips
            tblW.set(qn('w:type'), 'dxa')
            # Находим существующий tblW и заменяем или добавляем
            existing_tblW = tblPr.find(qn('w:tblW'))
            if existing_tblW is not None:
                tblPr.remove(existing_tblW)
            tblPr.insert(0, tblW)
            
            # Устанавливаем отрицательные поля слева через tblCellMar
            tblCellMar = OxmlElement('w:tblCellMar')
            left_margin = OxmlElement('w:left')
            left_margin.set(qn('w:w'), str(-750))  # -0.5 дюйма слева
            left_margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(left_margin)
            tblPr.append(tblCellMar)
            
            # Распределение ширины столбцов:
            # 1-й столбец - 7%
            # 3-й столбец (AdGroupName) - 8% (меньше)
            # Остальные 7 столбцов делят оставшиеся 85%
            table.columns[0].width = Inches(13.0 * 0.07)  # 7%
            table.columns[1].width = Inches(13.0 * 0.85 / 7)  # Один из 7 столбцов
            table.columns[2].width = Inches(13.0 * 0.08)  # 8% для AdGroupName
            for i in range(3, 9):
                table.columns[i].width = Inches(13.0 * 0.85 / 7)  # Остальные из 7 столбцов

            # Заголовки таблицы
            header_cells = table.rows[0].cells
            headers = [report_78, report_79, report_80, report_81, report_82, report_83, report_84, report_85, report_86]
            
            for i, header in enumerate(headers):
                paragraph = header_cells[i].paragraphs[0]
                paragraph.text = header
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                # Вертикальное выравнивание по центру
                header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Добавляем данные групп объявлений
            for i, adgroup in enumerate(adgroup_rows, 1):
                row = table.add_row()
                cells = row.cells
                
                # 1. Номер (сквозная нумерация)
                paragraph = cells[0].paragraphs[0]
                paragraph.text = str(i)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 2. Проверяем CampaignType
                campaign_type = adgroup.get('CampaignType', '')
                paragraph = cells[1].paragraphs[0]
                if campaign_type == 'TEXT_CAMPAIGN':
                    paragraph.text = report_77
                else:
                    paragraph.text = ''
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 3. AdGroupName
                paragraph = cells[2].paragraphs[0]
                paragraph.text = adgroup.get('AdGroupName', '')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 4. Всегда пишем report_87
                paragraph = cells[3].paragraphs[0]
                paragraph.text = report_87
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 5. Проверяем AdNetworkType
                ad_network_type = adgroup.get('AdNetworkType', '')
                paragraph = cells[4].paragraphs[0]
                if ad_network_type == 'SEARCH':
                    paragraph.text = report_88
                elif ad_network_type == 'AD_NETWORK':
                    paragraph.text = report_89
                else:
                    paragraph.text = ''
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 6. Clicks
                clicks = adgroup.get('Clicks', 0)
                paragraph = cells[5].paragraphs[0]
                paragraph.text = self.format_number_with_spaces(clicks)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 7. Cost (переводим из микрорублей в рубли и дополнительно делим на 10)
                cost = adgroup.get('Cost', 0)
                cost_in_rubles = cost / 1000000.0
                paragraph = cells[6].paragraphs[0]
                paragraph.text = self.format_number_with_two_decimals(cost_in_rubles)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 8. BounceRate
                bounce_rate = adgroup.get('BounceRate', 0)
                paragraph = cells[7].paragraphs[0]
                paragraph.text = self.format_number_with_spaces(bounce_rate)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 9. AvgCpc (переводим из микрорублей в рубли и дополнительно делим на 10)
                avg_cpc = adgroup.get('AvgCpc', 0)
                avg_cpc_in_rubles = avg_cpc / 1000000.0
                paragraph = cells[8].paragraphs[0]
                paragraph.text = self.format_number_with_two_decimals(avg_cpc_in_rubles)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                cells[8].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Добавляем итоговую строку из campaign_stats_summary
            total_row = table.add_row()
            cells = total_row.cells
            
            # Объединяем первые 5 ячеек для report_90
            paragraph = cells[0].paragraphs[0]
            paragraph.text = report_90
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Объединяем ячейки 0-4
            merged_cell = cells[0].merge(cells[4])
            
            # 6. Clicks из summary
            summary_clicks = campaign_stats.get('summary', {}).get('Clicks', 0)
            paragraph = cells[5].paragraphs[0]
            paragraph.text = self.format_number_with_spaces(summary_clicks)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 7. Cost из summary (переводим из микрорублей в рубли и дополнительно делим на 10)
            summary_cost = campaign_stats.get('summary', {}).get('Cost', 0)
            summary_cost_in_rubles = summary_cost / 1000000.0
            paragraph = cells[6].paragraphs[0]
            paragraph.text = self.format_number_with_two_decimals(summary_cost_in_rubles)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 8. BounceRate из summary
            summary_bounce_rate = campaign_stats.get('summary', {}).get('BounceRate', 0)
            paragraph = cells[7].paragraphs[0]
            paragraph.text = self.format_number_with_spaces(summary_bounce_rate)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 9. AvgCpc из summary (переводим из микрорублей в рубли и дополнительно делим на 10)
            summary_avg_cpc = campaign_stats.get('summary', {}).get('AvgCpc', 0)
            summary_avg_cpc_in_rubles = summary_avg_cpc / 1000000.0
            paragraph = cells[8].paragraphs[0]
            paragraph.text = self.format_number_with_two_decimals(summary_avg_cpc_in_rubles)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            cells[8].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 16. report_91 слева жирным
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(report_91)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        # 17. Добавляем все скриншоты из папки url_2
        try:
            # Формируем базовый путь к папке со скриншотами
            screenshots_folder = f"gen_report_context_contracts/data_yandex_direct/{report_id}_результаты/screenshots/url_2/"
            
            # Пробуем загрузить скриншоты (пытаемся найти все файлы screenshot_001.png, screenshot_002.png и т.д.)
            screenshot_index = 1
            while True:
                image_path = f"{screenshots_folder}screenshot_{screenshot_index:03d}.png"
                
                # Загружаем изображение из MinIO (silent=True, чтобы не выводить ошибку для несуществующих файлов)
                image_data = self.load_image_from_minio(image_path, silent=True)
                
                if image_data:
                    # Добавляем отступ перед скриншотом (начиная со второго)
                    if screenshot_index > 1:
                        doc.add_paragraph()
                    
                    # Создаем временный файл для изображения
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(image_data)
                        temp_file_path = temp_file.name
                    
                    try:
                        # Добавляем изображение в документ
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        # Вставляем изображение с максимальной шириной 6.0 дюймов (как для url_1)
                        run.add_picture(temp_file_path, width=Inches(6.0))
                            
                    finally:
                        # Удаляем временный файл
                        os.unlink(temp_file_path)
                    
                    screenshot_index += 1
                else:
                    # Если файл не найден, прекращаем попытки
                    if screenshot_index == 1:
                        print(f"❌ Не удалось загрузить скриншоты из папки: {screenshots_folder}")
                    break
                    
        except Exception as e:
            print(f"❌ Ошибка при добавлении скриншотов из url_2: {e}")
            raise e

    def create_section_11(self, doc: Document, report_data: Dict, report_id: int) -> None:
        """Создать одиннадцатый раздел отчета"""
        # Получаем тексты из БД
        report_92 = self.get_report_text('report_92')
        report_93 = self.get_report_text('report_93')
        report_94 = self.get_report_text('report_94')
        report_95 = self.get_report_text('report_95')
        report_96 = self.get_report_text('report_96')
        report_97 = self.get_report_text('report_97')
        report_98 = self.get_report_text('report_98')
        report_99 = self.get_report_text('report_99')
        report_100 = self.get_report_text('report_100')
        
        if not all([report_92, report_93, report_94, report_95, report_96, report_97, report_98, report_99, report_100]):
            print("❌ Не удалось получить все необходимые тексты для раздела 11")
            return
        
        # Загружаем данные из campaign_stats_summary
        campaign_stats = self.load_campaign_stats_summary(report_id)
        if not campaign_stats:
            print("❌ Не удалось загрузить данные campaign_stats_summary")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня: report_92 + date_request + № + application_number
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Форматируем дату в формате «17» марта 2025 г.
        date_request = report_data['date_request']
        if date_request:
            # Названия месяцев на русском в родительном падеже
            months_genitive = [
                'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
            ]
            day = date_request.day
            month = months_genitive[date_request.month - 1]
            year = date_request.year
            formatted_date = f'«{day}» {month} {year} г.'
        else:
            formatted_date = ""
        
        text = f"{report_92} {formatted_date} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_93 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_93} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 4. Пустая строка
        doc.add_paragraph()

        # 5. Добавляем все скриншоты из папки url_3
        try:
            # Формируем базовый путь к папке со скриншотами
            screenshots_folder = f"gen_report_context_contracts/data_yandex_direct/{report_id}_результаты/screenshots/url_3/"
            
            # Пробуем загрузить скриншоты (пытаемся найти все файлы screenshot_001.png, screenshot_002.png и т.д.)
            screenshot_index = 1
            while True:
                image_path = f"{screenshots_folder}screenshot_{screenshot_index:03d}.png"
                
                # Загружаем изображение из MinIO (silent=True, чтобы не выводить ошибку для несуществующих файлов)
                image_data = self.load_image_from_minio(image_path, silent=True)
                
                if image_data:
                    # Создаем временный файл для изображения
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(image_data)
                        temp_file_path = temp_file.name
                    
                    try:
                        # Добавляем изображение в документ
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        # Вставляем изображение с максимальной шириной 6.0 дюймов
                        run.add_picture(temp_file_path, width=Inches(6.0))
                        
                        # Добавляем небольшой отступ после каждого скриншота
                        doc.add_paragraph()
                            
                    finally:
                        # Удаляем временный файл
                        os.unlink(temp_file_path)
                    
                    screenshot_index += 1
                else:
                    # Если файл не найден, прекращаем попытки
                    if screenshot_index == 1:
                        print(f"❌ Не удалось загрузить скриншоты из папки: {screenshots_folder}")
                    break
                    
        except Exception as e:
            print(f"❌ Ошибка при добавлении скриншотов из url_3: {e}")
            raise e

        # 6. Пустая строка после скриншотов
        doc.add_paragraph()

        # 7. Создаём таблицу с данными
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'  # Стиль с видимыми границами
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы
        table.width = Inches(6.0)
        # Колонки: равномерно распределены по 50%
        table.columns[0].width = Inches(3.0)
        table.columns[1].width = Inches(3.0)
        
        # Строка 1: Объединенные ячейки с report_94 жирным
        cells = table.rows[0].cells
        merged_cell = cells[0].merge(cells[1])
        paragraph = merged_cell.paragraphs[0]
        paragraph.text = report_94
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Строка 2: report_95 | start_date - end_date
        cells = table.rows[1].cells
        
        # Первый столбец
        paragraph = cells[0].paragraphs[0]
        paragraph.text = report_95
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Второй столбец (жирный)
        start_date = report_data['start_date']
        end_date = report_data['end_date']
        if start_date and end_date:
            date_range = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
        else:
            date_range = ""
        paragraph = cells[1].paragraphs[0]
        paragraph.text = date_range
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Строка 3: report_96 | Clicks
        cells = table.rows[2].cells
        
        # Первый столбец
        paragraph = cells[0].paragraphs[0]
        paragraph.text = report_96
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Второй столбец (жирный)
        summary_clicks = campaign_stats.get('summary', {}).get('Clicks', 0)
        paragraph = cells[1].paragraphs[0]
        paragraph.text = self.format_number_with_spaces(summary_clicks)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Строка 4: report_97 | Cost + report_100
        cells = table.rows[3].cells
        
        # Первый столбец
        paragraph = cells[0].paragraphs[0]
        paragraph.text = report_97
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Второй столбец (жирный)
        summary_cost = campaign_stats.get('summary', {}).get('Cost', 0)
        cost_in_rubles = summary_cost / 1000000.0
        cost_text = f"{self.format_number_with_spaces(cost_in_rubles)} {report_100}"
        paragraph = cells[1].paragraphs[0]
        paragraph.text = cost_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Строка 5: report_98 | AvgCpc
        cells = table.rows[4].cells
        
        # Первый столбец
        paragraph = cells[0].paragraphs[0]
        paragraph.text = report_98
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Второй столбец (жирный)
        summary_avg_cpc = campaign_stats.get('summary', {}).get('AvgCpc', 0)
        avg_cpc_in_rubles = summary_avg_cpc / 1000000.0
        paragraph = cells[1].paragraphs[0]
        paragraph.text = self.format_number_with_spaces(avg_cpc_in_rubles)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Строка 6: report_99 | BounceRate + "%"
        cells = table.rows[5].cells
        
        # Первый столбец
        paragraph = cells[0].paragraphs[0]
        paragraph.text = report_99
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Второй столбец (жирный)
        summary_bounce_rate = campaign_stats.get('summary', {}).get('BounceRate', 0)
        bounce_rate_text = f"{self.format_number_with_spaces(summary_bounce_rate)}%"
        paragraph = cells[1].paragraphs[0]
        paragraph.text = bounce_rate_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def create_section_12(self, doc: Document, report_data: Dict, report_id: int) -> None:
        """Создать двенадцатый раздел отчета"""
        # Получаем тексты из БД
        report_101 = self.get_report_text('report_101')
        report_102 = self.get_report_text('report_102')
        
        if not all([report_101, report_102]):
            print("❌ Не удалось получить все необходимые тексты для раздела 12")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня: report_101 + date_request + № + application_number
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Форматируем дату в формате «17» марта 2025 г.
        date_request = report_data['date_request']
        if date_request:
            # Названия месяцев на русском в родительном падеже
            months_genitive = [
                'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
            ]
            day = date_request.day
            month = months_genitive[date_request.month - 1]
            year = date_request.year
            formatted_date = f'«{day}» {month} {year} г.'
        else:
            formatted_date = ""
        
        text = f"{report_101} {formatted_date} №{report_data['application_number']}"
        run = heading.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Тема контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_data['theme_contract'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Убираем интервал между абзацами и устанавливаем одинарный интервал
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 3. report_102 + дата и номер контракта
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Форматируем дату контракта в формате 28.12.2024 г.
        contract_date = report_data['date_contract']
        if contract_date:
            formatted_date = contract_date.strftime("%d.%m.%Y г.")
        else:
            formatted_date = ""
        text = f"{report_102} {formatted_date} №{report_data['number_contract']}"
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        # Устанавливаем одинарный интервал
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 4. Пустая строка
        doc.add_paragraph()

        # 5. Получаем тексты для таблицы объявлений
        report_103 = self.get_report_text('report_103')
        report_104 = self.get_report_text('report_104')
        report_105 = self.get_report_text('report_105')
        report_106 = self.get_report_text('report_106')
        report_107 = self.get_report_text('report_107')
        
        if not all([report_103, report_104, report_105, report_106, report_107]):
            print("❌ Не удалось получить все необходимые тексты для таблицы объявлений")
            return

        # 6. Получаем топ объявления
        top_ads = self.get_top_ads(report_id)
        if not top_ads:
            print("❌ Не удалось получить данные топ объявлений")
            return

        # 7. Создаем одну таблицу с лучшими объявлениями
        # Количество строк: 2 строки на каждое объявление (заголовок + данные)
        table_rows = len(top_ads) * 2
        table = doc.add_table(rows=table_rows, cols=2)
        table.style = 'Table Grid'
        table.allow_autofit = False
        
        # Настраиваем ширину таблицы
        table.width = Inches(6.5)
        table.columns[0].width = Inches(2.0)  # Изображение (еще уменьшено)
        table.columns[1].width = Inches(4.5)  # Данные (увеличено)
        
        # Заполняем таблицу данными
        for i, ad in enumerate(top_ads):
            row_index = i * 2
            
            # Строка 1: Объединенная ячейка с названием кампании
            cells = table.rows[row_index].cells
            merged_cell = cells[0].merge(cells[1])
            paragraph = merged_cell.paragraphs[0]
            paragraph.text = f"{report_103} {ad['CampaignName']}"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Строка 2: Изображение и данные
            cells = table.rows[row_index + 1].cells
            
            # Левая ячейка - изображение
            image_cell = cells[0]
            image_path = f"gen_report_context_contracts/data_yandex_direct/{report_id}_результаты/very_good_ads/{ad['AdId']}.png"
            
            try:
                # Загружаем изображение из MinIO
                image_data = self.load_image_from_minio(image_path, silent=True)
                
                if image_data:
                    # Создаем временный файл для изображения
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(image_data)
                        temp_file_path = temp_file.name
                    
                    try:
                        # Добавляем изображение в ячейку с ограничением по высоте
                        paragraph = image_cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        
                        # Определяем размеры изображения и масштабируем
                        with Image.open(temp_file_path) as img:
                            original_width, original_height = img.size
                            
                            # Максимальная высота 300px
                            max_height_px = 300
                            
                            # Рассчитываем пропорциональные размеры
                            if original_height > max_height_px:
                                # Масштабируем по высоте
                                scale_factor = max_height_px / original_height
                                new_width = int(original_width * scale_factor)
                                new_height = max_height_px
                                
                                # Конвертируем в дюймы
                                width_inches = Inches(new_width / 96)  # 96 DPI
                                height_inches = Inches(new_height / 96)
                            else:
                                # Изображение уже подходящего размера
                                width_inches = Inches(original_width / 96)
                                height_inches = Inches(original_height / 96)
                            
                            # Добавляем изображение с рассчитанными размерами
                            run.add_picture(temp_file_path, width=width_inches, height=height_inches)
                    finally:
                        # Удаляем временный файл
                        os.unlink(temp_file_path)
                else:
                    # Если изображение не найдено, выводим сообщение
                    paragraph = image_cell.paragraphs[0]
                    paragraph.text = "Не удалось получить скриншот объявления. Вставьте его вручную"
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    run = paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
                    
            except Exception as e:
                print(f"❌ Ошибка при загрузке изображения для объявления {ad['AdId']}: {e}")
                # Выводим сообщение об ошибке
                paragraph = image_cell.paragraphs[0]
                paragraph.text = "Не удалось получить скриншот изображения. Вставьте его вручную"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
            
            image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Правая ячейка - данные объявления
            data_cell = cells[1]
            data_paragraph = data_cell.paragraphs[0]
            
            # Форматируем AvgCpc (делим на 1000000)
            avg_cpc_rubles = ad['AvgCpc'] / 1000000.0 if ad['AvgCpc'] else 0
            
            # Создаем текст с данными (4 строки)
            data_text = f"{report_104} {ad['AdId']}\n"
            data_text += f"{report_105} {self.format_number_with_spaces(ad['Clicks'])}\n"
            data_text += f"{report_106} {self.format_number_with_spaces(ad['BounceRate'])}%\n"
            data_text += f"{report_107} {self.format_number_with_spaces(avg_cpc_rubles)}\n"
            data_text += f"{ad['Href']}"
            
            data_paragraph.text = data_text
            data_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            data_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = data_paragraph.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            data_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 8. Пустая строка после таблицы
        doc.add_paragraph()

        # 9. report_108 с новой строки
        report_108 = self.get_report_text('report_108')
        if report_108:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(report_108)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            # Устанавливаем одинарный интервал
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 10. Пустая строка
        doc.add_paragraph()

        # 11. Добавляем первые 10 скриншотов из папки url_4
        screenshots_folder = f"gen_report_context_contracts/data_yandex_direct/{report_id}_результаты/screenshots/url_4/"
        
        # Пробуем загрузить первые 10 скриншотов
        for screenshot_index in range(1, 11):  # от 1 до 10 включительно
            image_path = f"{screenshots_folder}screenshot_{screenshot_index:03d}.png"
            
            # Загружаем изображение из MinIO (silent=True, чтобы не выводить ошибку для несуществующих файлов)
            image_data = self.load_image_from_minio(image_path, silent=True)
            
            if image_data:
                # Создаем временный файл для изображения
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                    temp_file.write(image_data)
                    temp_file_path = temp_file.name
                
                try:
                    # Добавляем изображение в документ
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    # Вставляем изображение с максимальной шириной 6.0 дюймов
                    run.add_picture(temp_file_path, width=Inches(6.0))
                    
                    print(f"✅ Загружен скриншот: screenshot_{screenshot_index:03d}.png")
                    
                finally:
                    # Удаляем временный файл
                    os.unlink(temp_file_path)
            else:
                # Если файл не найден, это нормально - просто файлы закончились
                break

    def create_report(self, report_id: int) -> bool:
        """Создать отчет в формате Word"""
        try:
            # Получаем данные отчета
            report_data = self.get_report_data(report_id)
            if not report_data:
                print(f"❌ Не удалось получить данные отчета {report_id}")
                return False

            # Создаем новый документ
            doc = Document()

            # Настраиваем стиль по умолчанию
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # Настраиваем поля страницы (2.5 см со всех сторон)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Добавляем нумерацию страниц в правый нижний угол
            for section in sections:
                # Включаем разные колонтитулы для первой страницы
                section.different_first_page_header_footer = True
                
                # Основной колонтитул (для всех страниц кроме первой)
                footer = section.footer
                paragraph = footer.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Добавляем поле номера страницы
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                run = paragraph.add_run()
                run._r.append(fldChar1)
                
                instrText = OxmlElement('w:instrText')
                instrText.text = "PAGE"
                run._r.append(instrText)
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar2)
                
                # Колонтитул первой страницы остается пустым
                first_page_footer = section.first_page_footer
                if first_page_footer.paragraphs:
                    first_page_footer.paragraphs[0].clear()

            # Создаем разделы
            self.create_section_1(doc, report_data)  # Первый раздел
            self.create_section_2(doc, report_data)  # Таблица в первом разделе
            self.create_section_3(doc)  # Второй раздел с содержанием
            self.create_section_4(doc, report_data)  # Третий раздел с таблицей терминов
            self.create_section_5(doc, report_data)  # Четвертый раздел
            self.create_section_6(doc, report_data)  # Пятый раздел
            self.create_section_7(doc, report_data)  # Шестой раздел
            self.create_section_8(doc, report_data)  # Седьмой раздел
            self.create_section_9(doc, report_data, report_id)  # Восьмой раздел
            self.create_section_9_new(doc, report_data)  # Девятый раздел
            self.create_section_10(doc, report_data, report_id)  # Десятый раздел
            self.create_section_11(doc, report_data, report_id)  # Одиннадцатый раздел
            self.create_section_12(doc, report_data, report_id)  # Двенадцатый раздел
            self.create_section_13(doc, report_data, report_id)  # Тринадцатый раздел

            # Генерируем имя файла с датой и временем
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # filename = f"Отчет_{timestamp}.docx"
            filename = f"{report_id}/Отчет_{timestamp}.docx"
            output_path = os.path.join(self.output_folder, filename)

            # сохраняем файл
            # doc.save(output_path)

            # Сохраняем файл в S3
            file = io.BytesIO()
            doc.save(file)
            file.seek(0)
            # s3_file_path = os.getenv('S3_REPORT_PATH')
            # s3_file_path = '/'.join((s3_file_path, filename))
            # self.minio_client.put_object(self.bucket_name, s3_file_path, file, len(file.getvalue()))

            # записываем адрес (S3) в БД
            # write_s3path_to_bd(report_id, os.getenv('CONTENT_REPORT_COL_NAME'), s3_file_path)

            # print(f"✅ Word-файл создан: {output_path}")
            print(f"✅ Word-файл создан.")
            return file, filename
            
        except Exception as e:
            print(f"❌ Ошибка при создании Word-файла: {e}")
            import traceback
            traceback.print_exc()
            raise e

    def get_pending_reports(self) -> List[Dict]:
        """Получить отчеты со статусом 1 (готовые к обработке)"""
        conn = None
        try:
            print(f"🔌 Подключение к БД: {self.db_config['host']}:{self.db_config['port']}/{self.db_config['database']}")
            conn = self._connect_to_db()
            if not conn:
                return []
                
            cursor = conn.cursor()
            
            # Устанавливаем схему по умолчанию
            cursor.execute("SET search_path TO gen_report_context_contracts, public;")
            print("✓ Схема установлена: gen_report_context_contracts")
            
            # Получаем отчеты со статусом 1
            query = """
            SELECT 
                r.id, 
                r.id_contracts, 
                r.id_requests,
                c.number_contract,
                c.theme_contract,
                req.application_number
            FROM reports r
            JOIN contracts c ON r.id_contracts = c.id
            JOIN requests req ON r.id_requests = req.id
            WHERE r.id_status = 1 AND (r.is_deleted = false OR r.is_deleted IS NULL)
            ORDER BY r.create_entry DESC
            """
            
            print(f"🔍 Выполняем запрос...")
            cursor.execute(query)
            reports = []
            
            for row in cursor.fetchall():
                reports.append({
                    'id': row[0],
                    'id_contracts': row[1],
                    'id_requests': row[2],
                    'number_contract': row[3],
                    'theme_contract': row[4],
                    'application_number': row[5]
                })
            
            print(f"✅ Найдено отчетов для обработки: {len(reports)}")
            
            cursor.close()
            conn.close()
            
            return reports
            
        except Exception as e:
            print(f"❌ Ошибка при получении отчетов: {e}")
            if conn:
                conn.close()
            import traceback
            traceback.print_exc()
            # return []
            raise e

    def create_section_13(self, doc: Document, report_data: Dict, report_id: int) -> None:
        """Создать тринадцатый раздел отчета"""
        # Получаем текст из БД
        report_109 = self.get_report_text('report_109')
        conclusions_recommendations = report_data.get('conclusions_recommendations')
        
        if not all([report_109, conclusions_recommendations]):
            print("❌ Не удалось получить все необходимые тексты для раздела 13")
            return

        # Добавляем разрыв страницы
        doc.add_page_break()

        # 1. Заголовок первого уровня: report_109
        heading = doc.add_paragraph()
        heading.style = doc.styles['Heading 1']
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = heading.add_run(report_109)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Устанавливаем одинарный интервал
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 2. Пустая строка
        doc.add_paragraph()

        # 3. Получаем данные кампаний для поиска кампании с наибольшим количеством кликов
        try:
            # Загружаем данные кампаний
            campaign_stats_data = self.load_file_from_minio(report_id, f"campaign_stats_{report_id}.json")
            if not campaign_stats_data:
                print("❌ Не удалось загрузить данные кампаний")
                return

            campaigns_rows = campaign_stats_data.get('result', {}).get('rows', [])
            if not campaigns_rows:
                print("❌ Нет данных о кампаниях")
                return

            # Находим кампанию с наибольшим количеством кликов
            max_clicks_campaign = max(campaigns_rows, key=lambda x: x.get('Clicks', 0))
            
            # Извлекаем данные кампании
            campaign_id = max_clicks_campaign.get('CampaignId', '')
            campaign_name = max_clicks_campaign.get('CampaignName', '')
            clicks_count = max_clicks_campaign.get('Clicks', 0)
            
            # Определяем тип кампании из названия
            campaign_type = "РСЯ" if "РСЯ" in campaign_name else "Поиск"

            # Заменяем динамические параметры в тексте
            text_with_params = conclusions_recommendations
            text_with_params = text_with_params.replace('{number_campany}', str(campaign_id))
            text_with_params = text_with_params.replace('{title_campany}', campaign_name)
            text_with_params = text_with_params.replace('{type_campany}', campaign_type)
            text_with_params = text_with_params.replace('{number_clicks}', str(clicks_count))

            # Добавляем текст с подставленными параметрами
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text_with_params)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Устанавливаем полуторный интервал
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        except Exception as e:
            print(f"❌ Ошибка при обработке данных кампаний: {e}")
            # Добавляем текст без подстановки параметров
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(conclusions_recommendations)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Устанавливаем полуторный интервал
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def run(self):
        """Основной метод запуска обработки"""
        print("🚀 Запуск генератора отчетов...")
        
        try:
            # Получаем отчеты для обработки
            pending_reports = self.get_pending_reports()
            
            if not pending_reports:
                print("ℹ️ Нет отчетов для обработки")
                return
            
            # Обрабатываем каждый отчет
            for report_info in pending_reports:
                report_id = report_info['id']
                print(f"\n📋 Обработка отчета ID: {report_id}")
                print(f"   Контракт: {report_info['number_contract']}")
                print(f"   Тема: {report_info['theme_contract']}")
                print(f"   Номер заявки: {report_info['application_number']}")
                
                success = self.create_report(report_id)
                if success:
                    print(f"✅ Отчет {report_id} успешно создан")
                else:
                    print(f"❌ Не удалось создать отчет {report_id}")
                    
        except Exception as e:
            print(f"❌ Критическая ошибка: {e}")
            import traceback
            traceback.print_exc()
            raise e


def word_report_generate(report_id):
    """Главная функция"""
    try:
        generator = ReportGenerator()
        # Проверяем аргумент
        if report_id:
            # Если передан ID отчета как аргумент
            try:
                print(f"🎯 Обработка конкретного отчета ID: {report_id}")
                file, filename = generator.create_report(report_id)
                if file:
                    print(f"✅ Отчет {report_id} успешно создан")
                    return file, filename
                else:
                    print(f"❌ Не удалось создать отчет {report_id}")
            except ValueError:
                print("❌ Ошибка: ID отчета должен быть числом")
                print("Использование: python report_generator.py [ID_отчета]")
        else:
            # Обрабатываем все отчеты со статусом 1
            generator.run()
            
    except Exception as e:
        print(f"❌ Критическая ошибка: {e}")
        import traceback
        traceback.print_exc()
        raise e


if __name__ == "__main__":
    print(word_report_generate(16))
